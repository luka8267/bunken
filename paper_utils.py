import io
import math
import os
import re
import unicodedata
import uuid
import zipfile

from docx import Document

try:
    from citeproc import Citation, CitationItem, CitationStylesBibliography, CitationStylesStyle
    from citeproc.formatter import plain
    from citeproc.source.json import CiteProcJSON
    import citeproc_styles
except ImportError:
    Citation = None
    CitationItem = None
    CitationStylesBibliography = None
    CitationStylesStyle = None
    CiteProcJSON = None
    citeproc_styles = None
    plain = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from postgrest.exceptions import APIError
except ImportError:
    APIError = Exception

try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    genai = None
    genai_types = None

BUCKET_NAME = "paper-pdfs"
PAPER_ITEMS_VIEW = "paper_items_view"
PDF_SUMMARY_MAX_CHARS = 60000
ITEM_METADATA_COLUMNS = ("volume", "issue", "pages", "publisher", "item_type")
SAFE_STORAGE_NAME_RE = re.compile(r"[^A-Za-z0-9._-]+")
SAFE_STORAGE_EXT_RE = re.compile(r"[^A-Za-z0-9]")
DOI_RE = re.compile(r"10\.\d{4,9}/[^\s\"<>]+", re.IGNORECASE)
MAX_STORAGE_BASENAME_LENGTH = 80
READING_STATUSES = ["未読", "読書中", "読了", "再読したい", "引用予定"]
SORT_OPTIONS = ["追加順", "年（新しい順）", "年（古い順）", "タイトル", "ステータス"]
CSL_STYLE_OPTIONS = {
    "APA": "apa",
    "Vancouver": "vancouver",
    "Nature": "nature",
    "ACS": "american-chemical-society",
    "IEEE": "ieee",
    "Elsevier Harvard": "elsevier-harvard",
    "Chicago Author-Date": "chicago-author-date",
}
PDF_ANNOTATION_TYPES = {
    "highlight": "ハイライト",
    "page_note": "ページメモ",
    "citation_note": "引用予定",
}


def is_missing_relation_error(error):
    error_text = str(error).lower()
    return (
        "paper_items_view" in error_text
        or ("relation" in error_text and "does not exist" in error_text)
        or ("could not find the table" in error_text and "schema cache" in error_text)
        or ("could not find" in error_text and "relation" in error_text)
    )


def is_missing_metadata_column_error(error):
    error_text = str(error).lower()
    if "could not find" not in error_text and "column" not in error_text:
        return False
    return any(column in error_text for column in ITEM_METADATA_COLUMNS)


def strip_metadata_columns(columns):
    if columns == "*":
        return columns

    kept_columns = []
    for column in (columns or "").split(","):
        normalized = column.strip()
        if normalized and normalized not in ITEM_METADATA_COLUMNS:
            kept_columns.append(normalized)
    return ", ".join(kept_columns) or columns


def is_duplicate_key_error(error):
    error_text = str(error).lower()
    return (
        "duplicate" in error_text
        or "23505" in error_text
        or "already exists" in error_text
    )


def is_permission_error(error):
    error_text = str(error).lower()
    return (
        "42501" in error_text
        or "permission denied" in error_text
        or "row-level security" in error_text
        or "violates row-level security" in error_text
    )


def normalize_json_value(value):
    if value is None:
        return None
    if value.__class__.__name__ in {"NAType", "NaTType"}:
        return None
    try:
        if value != value:
            return None
    except Exception:
        pass
    if isinstance(value, float) and not math.isfinite(value):
        return None
    return value


def normalize_optional_db_value(value):
    value = normalize_json_value(value)
    if value is None:
        return None
    if isinstance(value, str):
        return value or None
    return value


def has_attachment_path(value):
    value = normalize_json_value(value)
    if value is None:
        return False
    text = str(value).strip()
    if not text:
        return False
    return text.casefold() not in {"none", "null", "nan", "na", "n/a", "[]", "{}"}


def normalize_text_db_value(value):
    value = normalize_json_value(value)
    if value is None:
        return ""
    return str(value)


def normalize_doi(doi):
    doi = normalize_json_value(doi)
    if doi is None:
        return ""
    text = str(doi).strip()
    if not isinstance(doi, str):
        match = DOI_RE.search(text)
        return match.group(0).rstrip(").,;]") if match else ""
    if not text:
        return ""
    text = re.sub(r"^doi:\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", text, flags=re.IGNORECASE)
    if text.lower().startswith(("http://", "https://")):
        match = DOI_RE.search(text)
        if match:
            text = match.group(0)
    text = text.strip().strip("<>").rstrip(").,;]")
    return text


def extract_doi_from_text(text):
    match = DOI_RE.search(text or "")
    if not match:
        return ""
    return match.group(0).rstrip(").,;]")


def normalize_title_for_match(title):
    normalized = unicodedata.normalize("NFKC", title or "").casefold()
    return re.sub(r"\s+", " ", normalized).strip()


JOURNAL_NORMALIZATION_MAP = {
    "chem eur j": "Chemistry - A European Journal",
    "chemistry a european journal": "Chemistry - A European Journal",
    "j am chem soc": "Journal of the American Chemical Society",
    "journal of the american chemical society": "Journal of the American Chemical Society",
    "angew chem int ed": "Angewandte Chemie International Edition",
    "angewandte chemie international edition": "Angewandte Chemie International Edition",
    "phys chem chem phys": "Physical Chemistry Chemical Physics",
    "physical chemistry chemical physics": "Physical Chemistry Chemical Physics",
}


def normalize_author_name(name):
    text = re.sub(r"\s+", " ", unicodedata.normalize("NFKC", name or "")).strip()
    if not text:
        return ""
    if "," in text:
        family, given = [part.strip() for part in text.split(",", maxsplit=1)]
        return f"{family}, {given}" if given else family
    parts = text.split()
    if len(parts) >= 2:
        return f"{parts[-1]}, {' '.join(parts[:-1])}"
    return text


def normalize_author_list(authors):
    return ", ".join(
        name
        for name in (
            normalize_author_name(part)
            for part in re.split(r"\s+and\s+|;|\|", authors or "")
        )
        if name
    )


def normalize_journal_title(journal):
    text = re.sub(r"\s+", " ", unicodedata.normalize("NFKC", journal or "")).strip()
    if not text:
        return ""
    key = re.sub(r"[^a-z0-9]+", " ", text.casefold()).strip()
    return JOURNAL_NORMALIZATION_MAP.get(key, text)


def find_duplicate_paper_groups(papers):
    groups_by_key = {}

    for paper in papers or []:
        doi = normalize_doi(paper.get("doi")).lower()
        if doi:
            groups_by_key.setdefault(("DOI", doi), []).append(paper)

        title = normalize_title_for_match(paper.get("title"))
        year = paper.get("year")
        if title and year:
            groups_by_key.setdefault(("タイトル+年", f"{title}:{year}"), []).append(paper)

        if title:
            groups_by_key.setdefault(("タイトル類似", title), []).append(paper)

        author_key = normalize_author_list(paper.get("authors")).casefold()
        journal_key = normalize_journal_title(paper.get("journal")).casefold()
        if title and author_key:
            groups_by_key.setdefault(("タイトル+著者", f"{title}:{author_key}"), []).append(paper)
        if title and journal_key and year:
            groups_by_key.setdefault(("タイトル+雑誌+年", f"{title}:{journal_key}:{year}"), []).append(paper)

    duplicate_groups = []
    for (reason, value), group_papers in groups_by_key.items():
        if len(group_papers) < 2:
            continue

        duplicate_groups.append(
            {
                "reason": reason,
                "value": value,
                "papers": group_papers,
            }
        )

    return duplicate_groups


def describe_duplicate_group(group):
    reason = group.get("reason") or ""
    papers = group.get("papers") or []
    dois = {
        normalize_doi(paper.get("doi")).casefold()
        for paper in papers
        if normalize_doi(paper.get("doi"))
    }
    titles = {
        normalize_title_for_match(paper.get("title"))
        for paper in papers
        if normalize_title_for_match(paper.get("title"))
    }
    years = {
        str(paper.get("year") or "").strip()
        for paper in papers
        if str(paper.get("year") or "").strip()
    }
    authors = {
        normalize_author_list(paper.get("authors")).casefold()
        for paper in papers
        if normalize_author_list(paper.get("authors"))
    }
    journals = {
        normalize_journal_title(paper.get("journal")).casefold()
        for paper in papers
        if normalize_journal_title(paper.get("journal"))
    }

    evidence = []
    score = 0
    if reason == "DOI" or (len(dois) == 1 and dois):
        score += 80
        evidence.append("DOIが一致")
    if len(titles) == 1 and titles:
        score += 25
        evidence.append("タイトルが一致")
    if len(authors) == 1 and authors:
        score += 15
        evidence.append("著者が一致")
    if len(journals) == 1 and journals:
        score += 10
        evidence.append("雑誌名が一致")
    if len(years) == 1 and years:
        score += 10
        evidence.append("年が一致")
    elif len(years) > 1:
        score -= 15
        evidence.append("年が異なる")

    if reason == "タイトル類似" and not dois:
        score = min(score, 55)
        evidence.append("DOIなしのタイトル一致")

    score = max(0, min(score, 100))
    if score >= 80:
        level = "強い候補"
        advice = "同じ文献である可能性が高いです。残す文献と補完する値を確認して統合してください。"
    elif score >= 55:
        level = "中程度"
        advice = "同じ文献の可能性があります。著者、年、雑誌名を確認してから統合してください。"
    else:
        level = "要確認"
        advice = "誤統合の可能性があります。内容をよく確認し、迷う場合は統合しないでください。"

    return {
        "score": score,
        "level": level,
        "evidence": list(dict.fromkeys(evidence)) or [reason or "条件一致"],
        "advice": advice,
    }


def normalize_tag_input(tags_text):
    seen = set()
    normalized = []
    for tag in (tags_text or "").split(","):
        value = tag.strip()
        if value and value not in seen:
            seen.add(value)
            normalized.append(value)
    return normalized


def is_storage_path(value):
    return isinstance(value, str) and bool(value.strip())


def normalize_optional_id(value):
    if value is None:
        return None
    if value != value:
        return None
    text = str(value).strip()
    return text or None


def normalize_uuid_text(value):
    text = normalize_optional_id(value)
    if not text:
        return None
    try:
        return str(uuid.UUID(text))
    except (TypeError, ValueError):
        return None


def make_safe_storage_filename(filename, default_ext=".pdf"):
    name, ext = os.path.splitext(filename or "")
    ext = ext.lower()
    if ext:
        ext = f".{SAFE_STORAGE_EXT_RE.sub('', ext.lstrip('.'))}"
    ext = ext if ext and ext != "." else default_ext

    normalized_name = unicodedata.normalize("NFKD", name)
    ascii_name = normalized_name.encode("ascii", "ignore").decode("ascii")
    safe_name = SAFE_STORAGE_NAME_RE.sub("-", ascii_name).strip(".-_")
    safe_name = safe_name[:MAX_STORAGE_BASENAME_LENGTH].strip(".-_") or "paper"

    return f"{uuid.uuid4().hex}_{safe_name}{ext}"


def fetch_user_papers(supabase, user_id, columns="*"):
    try:
        return (
            supabase.table(PAPER_ITEMS_VIEW)
            .select(columns)
            .eq("user_id", user_id)
            .order("display_order")
            .execute()
        )
    except APIError as error:
        if is_missing_metadata_column_error(error):
            fallback_columns = strip_metadata_columns(columns)
            if fallback_columns != columns:
                return fetch_user_papers(supabase, user_id, fallback_columns)
        if is_missing_relation_error(error):
            return (
                supabase.table("papers")
                .select(strip_metadata_columns(columns))
                .eq("user_id", user_id)
                .order("display_order")
                .execute()
            )
        error_text = str(error).lower()
        if "uuid" in error_text or "user_id" in error_text:
            raise RuntimeError(
                "papers.user_id と認証ユーザーの紐づけ、または RLS 設定を確認してください。"
            ) from error
        raise


def search_user_papers(supabase, user_id, keyword, columns="id, title, authors, year"):
    normalized_keyword = (keyword or "").strip()
    query = (
        supabase.table(PAPER_ITEMS_VIEW)
        .select(columns)
        .eq("user_id", user_id)
        .order("display_order")
    )

    if normalized_keyword:
        escaped_keyword = normalized_keyword.replace("%", "\\%").replace(",", "\\,")
        query = query.or_(
            f"title.ilike.%{escaped_keyword}%,authors.ilike.%{escaped_keyword}%"
        )

    try:
        return query.execute()
    except APIError as error:
        if is_missing_metadata_column_error(error):
            fallback_columns = strip_metadata_columns(columns)
            if fallback_columns != columns:
                return search_user_papers(
                    supabase,
                    user_id,
                    keyword,
                    fallback_columns,
                )
        if not is_missing_relation_error(error):
            raise

    query = (
        supabase.table("papers")
        .select(strip_metadata_columns(columns))
        .eq("user_id", user_id)
        .order("display_order")
    )

    if normalized_keyword:
        escaped_keyword = normalized_keyword.replace("%", "\\%").replace(",", "\\,")
        query = query.or_(
            f"title.ilike.%{escaped_keyword}%,authors.ilike.%{escaped_keyword}%"
        )

    return query.execute()


def filter_papers(
    papers,
    keyword="",
    year_from=None,
    year_to=None,
    status="",
    attachment_filter="",
):
    normalized_keyword = (keyword or "").casefold().strip()
    filtered = []

    for paper in papers or []:
        if normalized_keyword:
            haystack = " ".join(
                str(paper.get(field) or "")
                for field in ("title", "authors", "journal", "doi", "notes")
            ).casefold()
            if normalized_keyword not in haystack:
                continue

        year = paper.get("year")
        if year_from is not None and year and int(year) < year_from:
            continue
        if year_to is not None and year and int(year) > year_to:
            continue

        if status and paper.get("status") != status:
            continue

        has_pdf = has_attachment_path(paper.get("pdf_path"))
        has_supporting = has_attachment_path(paper.get("supporting_path"))
        if attachment_filter == "PDFあり" and not has_pdf:
            continue
        if attachment_filter == "補足資料あり" and not has_supporting:
            continue
        if attachment_filter == "添付あり" and not (has_pdf or has_supporting):
            continue
        if attachment_filter == "添付なし" and (has_pdf or has_supporting):
            continue

        filtered.append(paper)

    return filtered


def fetch_user_papers_by_ids(supabase, user_id, paper_ids, columns="id, title"):
    if not paper_ids:
        return None

    try:
        return (
            supabase.table(PAPER_ITEMS_VIEW)
            .select(columns)
            .eq("user_id", user_id)
            .in_("id", paper_ids)
            .execute()
        )
    except APIError as error:
        if is_missing_metadata_column_error(error):
            fallback_columns = strip_metadata_columns(columns)
            if fallback_columns != columns:
                return fetch_user_papers_by_ids(
                    supabase,
                    user_id,
                    paper_ids,
                    fallback_columns,
                )
        if not is_missing_relation_error(error):
            raise

    return (
        supabase.table("papers")
        .select(strip_metadata_columns(columns))
        .eq("user_id", user_id)
        .in_("id", paper_ids)
        .execute()
    )


def find_existing_user_paper_by_doi(supabase, user_id, doi, columns="id, title"):
    normalized_doi = normalize_doi(doi)
    if not normalized_doi:
        return None

    try:
        result = (
            supabase.table(PAPER_ITEMS_VIEW)
            .select(columns)
            .eq("user_id", user_id)
            .eq("doi", normalized_doi)
            .limit(1)
            .execute()
        )
    except APIError as error:
        if is_missing_metadata_column_error(error):
            fallback_columns = strip_metadata_columns(columns)
            if fallback_columns != columns:
                return find_existing_user_paper_by_doi(
                    supabase,
                    user_id,
                    doi,
                    fallback_columns,
                )
        if not is_missing_relation_error(error):
            raise
        result = (
            supabase.table("papers")
            .select(strip_metadata_columns(columns))
            .eq("user_id", user_id)
            .eq("doi", normalized_doi)
            .limit(1)
            .execute()
        )

    return (result.data or [None])[0]


def get_next_display_order(supabase, user_id):
    try:
        result = fetch_user_papers(
            supabase,
            user_id,
            columns="display_order",
        )
        values = [
            int(row["display_order"])
            for row in (result.data or [])
            if row.get("display_order") is not None
        ]
        return (max(values) if values else 0) + 1
    except Exception:
        result = (
            supabase.table("papers")
            .select("display_order")
            .eq("user_id", user_id)
            .order("display_order", desc=True)
            .limit(1)
            .execute()
        )
        current_max = result.data[0]["display_order"] if result.data else 0
        return (current_max or 0) + 1


def create_attachment_row(supabase, user_id, item_id, kind, storage_path):
    if not is_storage_path(storage_path):
        return

    (
        supabase.table("attachments")
        .insert(
            {
                "item_id": item_id,
                "user_id": user_id,
                "kind": kind,
                "storage_path": storage_path,
            }
        )
        .execute()
    )


def create_item_creator_rows(supabase, item_id, authors):
    names = [name.strip() for name in (authors or "").split(",") if name.strip()]
    for position, name in enumerate(names, start=1):
        (
            supabase.table("creators")
            .insert(
                {
                    "item_id": item_id,
                    "creator_type": "author",
                    "literal_name": name,
                    "position": position,
                }
            )
            .execute()
        )


def create_item_backed_paper(
    supabase,
    user_id,
    title,
    authors,
    journal,
    year,
    doi,
    url,
    pdf_path,
    supporting_path,
    status,
    notes,
    display_order,
    volume="",
    issue="",
    pages="",
    publisher="",
    item_type="journalArticle",
):
    item_payload = {
        "user_id": user_id,
        "item_type": item_type or "journalArticle",
        "title": title,
        "publication_title": journal,
        "year": int(year),
        "doi": doi or None,
        "url": url or None,
        "abstract_note": notes,
        "extra": {
            "legacy_status": status,
            "legacy_display_order": str(display_order),
        },
    }
    for field, value in (
        ("volume", volume),
        ("issue", issue),
        ("pages", pages),
        ("publisher", publisher),
    ):
        if value:
            item_payload[field] = value

    item_result = (
        supabase.table("items")
        .insert(item_payload)
        .execute()
    )
    item_id = item_result.data[0]["id"]
    create_item_creator_rows(supabase, item_id, authors)
    create_attachment_row(supabase, user_id, item_id, "pdf", pdf_path)
    create_attachment_row(supabase, user_id, item_id, "supporting", supporting_path)
    return {"id": str(item_id), "item_id": item_id}


def create_legacy_paper(
    supabase,
    user_id,
    title,
    authors,
    journal,
    year,
    doi,
    url,
    pdf_path,
    supporting_path,
    status,
    notes,
    display_order,
    volume="",
    issue="",
    pages="",
    publisher="",
    item_type="journalArticle",
):
    insert_result = (
        supabase.table("papers")
        .insert(
            {
                "title": title,
                "authors": authors,
                "journal": journal,
                "year": int(year),
                "doi": doi or None,
                "url": url or None,
                "pdf_path": pdf_path,
                "supporting_path": supporting_path,
                "user_id": user_id,
                "display_order": display_order,
                "status": status,
                "notes": notes,
            }
        )
        .execute()
    )
    return {"id": insert_result.data[0]["id"], "item_id": None}


def create_user_paper(
    supabase,
    user_id,
    title,
    authors,
    journal,
    year,
    doi,
    url,
    pdf_path,
    supporting_path,
    status,
    notes,
    display_order,
    volume="",
    issue="",
    pages="",
    publisher="",
    item_type="journalArticle",
):
    try:
        try:
            return create_item_backed_paper(
                supabase,
                user_id,
                title,
                authors,
                journal,
                year,
                doi,
                url,
                pdf_path,
                supporting_path,
                status,
                notes,
                display_order,
                volume,
                issue,
                pages,
                publisher,
                item_type,
            )
        except APIError as error:
            if not is_missing_metadata_column_error(error):
                raise
            return create_item_backed_paper(
                supabase,
                user_id,
                title,
                authors,
                journal,
                year,
                doi,
                url,
                pdf_path,
                supporting_path,
                status,
                notes,
                display_order,
            )
    except APIError as error:
        if not is_missing_relation_error(error) and "items" not in str(error).lower():
            raise

    return create_legacy_paper(
        supabase,
        user_id,
        title,
        authors,
        journal,
        year,
        doi,
        url,
        pdf_path,
        supporting_path,
        status,
        notes,
        display_order,
        volume,
        issue,
        pages,
        publisher,
        item_type,
    )


def fetch_user_documents(supabase, user_id):
    return (
        supabase.table("documents")
        .select("id, word_document_id, title, citation_style, locale, updated_at")
        .eq("user_id", user_id)
        .order("updated_at", desc=True)
        .execute()
    )


def update_user_document_style(supabase, user_id, document_id, citation_style, locale=None):
    fields = {"citation_style": normalize_optional_db_value(citation_style) or "vancouver"}
    if locale is not None:
        fields["locale"] = normalize_optional_db_value(locale)
    return (
        supabase.table("documents")
        .update(fields)
        .eq("id", document_id)
        .eq("user_id", user_id)
        .execute()
    )


def fetch_document_citations(supabase, document_id):
    return (
        supabase.table("document_citations")
        .select(
            "id, citation_key, word_control_id, citation_items, rendered_text, "
            "context_text, sort_order, updated_at"
        )
        .eq("document_id", document_id)
        .order("sort_order")
        .execute()
    )


def delete_user_document(supabase, user_id, document_id):
    (
        supabase.table("document_citations")
        .delete()
        .eq("document_id", document_id)
        .execute()
    )
    return (
        supabase.table("documents")
        .delete()
        .eq("id", document_id)
        .eq("user_id", user_id)
        .execute()
    )


def citation_search_text(citation, paper_map):
    parts = [
        citation.get("rendered_text"),
        citation.get("context_text"),
        citation.get("updated_at"),
    ]
    for item in citation.get("citation_items") or []:
        if not isinstance(item, dict):
            continue
        parts.extend(
            [
                item.get("locator"),
                item.get("referenceNumber"),
            ]
        )
        paper = paper_map.get(str(item.get("paperId") or ""))
        if paper:
            parts.extend(
                [
                    paper.get("title"),
                    paper.get("authors"),
                    paper.get("journal"),
                    paper.get("year"),
                    paper.get("doi"),
                ]
            )
    return " ".join(str(part) for part in parts if part not in (None, ""))


def filter_document_citations(citations, paper_map, keyword):
    normalized_keyword = (keyword or "").strip().lower()
    if not normalized_keyword:
        return list(citations)
    return [
        citation
        for citation in citations
        if normalized_keyword in citation_search_text(citation, paper_map).lower()
    ]


def build_document_citation_export_rows(citations, paper_map):
    rows = []
    for citation in citations:
        base = {
            "順序": citation.get("sort_order") or "",
            "引用表示": citation.get("rendered_text") or "",
            "引用に使った文": citation.get("context_text") or "",
            "更新日時": citation.get("updated_at") or "",
        }
        items = [
            item
            for item in (citation.get("citation_items") or [])
            if isinstance(item, dict)
        ]
        if not items:
            rows.append(
                {
                    **base,
                    "文献タイトル": "",
                    "著者": "",
                    "年": "",
                    "掲載誌": "",
                    "DOI": "",
                    "参考文献番号": "",
                    "位置": "",
                }
            )
            continue

        for item in items:
            paper = paper_map.get(str(item.get("paperId") or "")) or {}
            rows.append(
                {
                    **base,
                    "文献タイトル": paper.get("title") or "",
                    "著者": paper.get("authors") or "",
                    "年": paper.get("year") or "",
                    "掲載誌": paper.get("journal") or "",
                    "DOI": paper.get("doi") or "",
                    "参考文献番号": item.get("referenceNumber") or "",
                    "位置": item.get("locator") or "",
                }
            )
    return rows


def get_document_citation_usage_map(supabase, user_id, papers):
    paper_refs = {
        str(reference_id)
        for paper in papers
        for reference_id in get_paper_reference_ids(paper)
    }
    usage_map = {reference_id: [] for reference_id in paper_refs}
    if not paper_refs:
        return usage_map

    documents_result = fetch_user_documents(supabase, user_id)
    for document in documents_result.data or []:
        citations_result = fetch_document_citations(supabase, document["id"])
        for citation in citations_result.data or []:
            for item in citation.get("citation_items") or []:
                if not isinstance(item, dict):
                    continue
                paper_id = str(item.get("paperId") or "")
                if paper_id not in paper_refs:
                    continue
                usage_map.setdefault(paper_id, []).append(
                    {
                        "document_title": document.get("title") or "無題",
                        "citation_text": citation.get("rendered_text") or "",
                        "context_text": citation.get("context_text") or "",
                        "reference_number": item.get("referenceNumber"),
                        "locator": item.get("locator"),
                        "updated_at": citation.get("updated_at") or "",
                    }
                )

    return usage_map


def replace_paper_id_in_document_citations(supabase, user_id, source_paper_id, target_paper_id):
    source_ids = source_paper_id if isinstance(source_paper_id, (list, tuple, set)) else [source_paper_id]
    source_texts = {str(source_id) for source_id in source_ids if source_id is not None}
    updated_count = 0

    documents_result = fetch_user_documents(supabase, user_id)
    for document in documents_result.data or []:
        citations_result = fetch_document_citations(supabase, document["id"])
        for citation in citations_result.data or []:
            citation_items = citation.get("citation_items") or []
            changed = False
            for item in citation_items:
                if not isinstance(item, dict):
                    continue
                if str(item.get("paperId")) in source_texts:
                    item["paperId"] = target_paper_id
                    changed = True

            if changed:
                (
                    supabase.table("document_citations")
                    .update({"citation_items": citation_items})
                    .eq("id", citation["id"])
                    .execute()
                )
                updated_count += 1

    return updated_count


def paper_has_document_citation_refs(supabase, user_id, paper_id):
    paper_ids = paper_id if isinstance(paper_id, (list, tuple, set)) else [paper_id]
    paper_texts = {str(value) for value in paper_ids if value is not None}
    documents_result = fetch_user_documents(supabase, user_id)
    for document in documents_result.data or []:
        citations_result = fetch_document_citations(supabase, document["id"])
        for citation in citations_result.data or []:
            for item in citation.get("citation_items") or []:
                if not isinstance(item, dict):
                    continue
                if str(item.get("paperId")) in paper_texts:
                    return True
    return False


def fetch_user_collections(supabase, user_id):
    return (
        supabase.table("collections")
        .select("id, name, parent_id, sort_order")
        .eq("user_id", user_id)
        .order("sort_order")
        .order("name")
        .execute()
    )


def create_collection(supabase, user_id, name):
    normalized_name = (name or "").strip()
    if not normalized_name:
        raise ValueError("Collection name is required.")

    max_result = (
        supabase.table("collections")
        .select("sort_order")
        .eq("user_id", user_id)
        .order("sort_order", desc=True)
        .limit(1)
        .execute()
    )
    current_max = max_result.data[0]["sort_order"] if max_result.data else 0

    return (
        supabase.table("collections")
        .insert(
            {
                "user_id": user_id,
                "name": normalized_name,
                "sort_order": (current_max or 0) + 1,
            }
        )
        .execute()
    )


def update_collection(supabase, user_id, collection_id, name):
    normalized_name = (name or "").strip()
    if not normalized_name:
        raise ValueError("Collection name is required.")

    return (
        supabase.table("collections")
        .update({"name": normalized_name})
        .eq("id", collection_id)
        .eq("user_id", user_id)
        .execute()
    )


def delete_collection(supabase, user_id, collection_id):
    return (
        supabase.table("collections")
        .delete()
        .eq("id", collection_id)
        .eq("user_id", user_id)
        .execute()
    )


def fetch_collection_paper_ids(supabase, collection_id):
    result = (
        supabase.table("collection_papers")
        .select("paper_id")
        .eq("collection_id", collection_id)
        .execute()
    )
    return [str(row["paper_id"]) for row in (result.data or [])]


def fetch_collection_item_ids(supabase, collection_id):
    try:
        result = (
            supabase.table("collection_items")
            .select("item_id")
            .eq("collection_id", collection_id)
            .execute()
        )
    except APIError as error:
        if is_missing_relation_error(error):
            return []
        raise
    return [str(row["item_id"]) for row in (result.data or [])]


def fetch_collection_counts(supabase, collection_ids):
    if not collection_ids:
        return {}

    references_by_collection = {collection_id: set() for collection_id in collection_ids}

    paper_result = (
        supabase.table("collection_papers")
        .select("collection_id, paper_id")
        .in_("collection_id", collection_ids)
        .execute()
    )
    for row in paper_result.data or []:
        collection_id = row.get("collection_id")
        paper_id = normalize_optional_id(row.get("paper_id"))
        if collection_id in references_by_collection and paper_id:
            references_by_collection[collection_id].add(("paper", paper_id))

    try:
        item_result = (
            supabase.table("collection_items")
            .select("collection_id, item_id")
            .in_("collection_id", collection_ids)
            .execute()
        )
    except APIError as error:
        if is_missing_relation_error(error):
            return {
                collection_id: len(references)
                for collection_id, references in references_by_collection.items()
            }
        raise

    item_rows = item_result.data or []
    item_ids = [
        item_id
        for item_id in (normalize_optional_id(row.get("item_id")) for row in item_rows)
        if item_id
    ]
    legacy_key_by_item_id = {}
    if item_ids:
        item_metadata = (
            supabase.table("items")
            .select("id, legacy_source, legacy_paper_id")
            .in_("id", item_ids)
            .execute()
        )
        for row in item_metadata.data or []:
            item_id = normalize_optional_id(row.get("id"))
            legacy_paper_id = normalize_optional_id(row.get("legacy_paper_id"))
            if (
                item_id
                and row.get("legacy_source") == "papers"
                and legacy_paper_id
            ):
                legacy_key_by_item_id[item_id] = ("paper", legacy_paper_id)

    for row in item_rows:
        collection_id = row.get("collection_id")
        item_id = normalize_optional_id(row.get("item_id"))
        if collection_id in references_by_collection and item_id:
            references_by_collection[collection_id].add(
                legacy_key_by_item_id.get(item_id, ("item", item_id))
            )

    return {
        collection_id: len(references)
        for collection_id, references in references_by_collection.items()
    }


def fetch_paper_collection_ids(supabase, paper_id, item_id=None):
    collection_ids = set()
    if paper_id and not normalize_uuid_text(paper_id):
        paper_result = (
            supabase.table("collection_papers")
            .select("collection_id")
            .eq("paper_id", paper_id)
            .execute()
        )
        collection_ids.update(row["collection_id"] for row in (paper_result.data or []))

    if item_id:
        try:
            item_result = (
                supabase.table("collection_items")
                .select("collection_id")
                .eq("item_id", item_id)
                .execute()
            )
            collection_ids.update(row["collection_id"] for row in (item_result.data or []))
        except APIError as error:
            if not is_missing_relation_error(error):
                raise

    return sorted(collection_ids)


def add_column_if_missing(columns, column):
    if columns == "*":
        return columns
    selected_columns = [value.strip() for value in columns.split(",")]
    if column in selected_columns:
        return columns
    return f"{columns}, {column}"


def fetch_papers_for_collection(supabase, user_id, collection_id, columns="id, title, authors, year"):
    paper_ids = fetch_collection_paper_ids(supabase, collection_id)
    item_ids = fetch_collection_item_ids(supabase, collection_id)
    reference_ids = set(paper_ids + item_ids)
    if not reference_ids:
        return []

    fetch_columns = add_column_if_missing(columns, "item_id") if item_ids else columns
    try:
        result = fetch_user_papers(supabase, user_id, columns=fetch_columns)
    except APIError:
        result = fetch_user_papers(supabase, user_id, columns=columns)
    return [
        paper
        for paper in (result.data or [])
        if str(paper.get("id")) in reference_ids
        or str(paper.get("item_id")) in reference_ids
    ]


def set_paper_collections(supabase, paper_id, selected_collection_ids, item_id=None):
    desired_ids = set(selected_collection_ids or [])
    current_ids = set(fetch_paper_collection_ids(supabase, paper_id, item_id))
    table_name = "collection_items" if item_id else "collection_papers"
    id_column = "item_id" if item_id else "paper_id"
    record_id = item_id or paper_id
    item_collection_table_missing = False
    legacy_paper_id = paper_id if paper_id and not normalize_uuid_text(paper_id) else None

    for collection_id in sorted(current_ids - desired_ids):
        if item_id and legacy_paper_id:
            (
                supabase.table("collection_papers")
                .delete()
                .eq("paper_id", legacy_paper_id)
                .eq("collection_id", collection_id)
                .execute()
            )
        try:
            (
                supabase.table(table_name)
                .delete()
                .eq(id_column, record_id)
                .eq("collection_id", collection_id)
                .execute()
            )
        except APIError as error:
            if not (
                item_id
                and (is_missing_relation_error(error) or is_permission_error(error))
            ):
                raise
            item_collection_table_missing = True

    for collection_id in sorted(desired_ids - current_ids):
        try:
            (
                supabase.table(table_name)
                .insert({id_column: record_id, "collection_id": collection_id})
                .execute()
            )
        except APIError as error:
            if item_id and (is_missing_relation_error(error) or is_permission_error(error)):
                item_collection_table_missing = True
                if not legacy_paper_id:
                    raise
                try:
                    (
                        supabase.table("collection_papers")
                        .insert(
                            {
                                "paper_id": legacy_paper_id,
                                "collection_id": collection_id,
                            }
                        )
                        .execute()
                    )
                except APIError as fallback_error:
                    if not is_duplicate_key_error(fallback_error):
                        raise RuntimeError(
                            "collection_items migration is not applied, and the "
                            "legacy collection_papers fallback could not save this "
                            "membership. Apply the normalized collection migration."
                        ) from fallback_error
            elif not is_duplicate_key_error(error):
                raise

    if item_collection_table_missing:
        return


def copy_paper_tags(supabase, source_paper_id, target_paper_id):
    tag_result = (
        supabase.table("paper_tags")
        .select("tag_id")
        .eq("paper_id", source_paper_id)
        .execute()
    )

    for row in tag_result.data or []:
        try:
            (
                supabase.table("paper_tags")
                .insert({"paper_id": target_paper_id, "tag_id": row["tag_id"]})
                .execute()
            )
        except APIError as error:
            if not is_duplicate_key_error(error):
                raise


def copy_paper_collections(supabase, source_paper_id, target_paper_id):
    collection_result = (
        supabase.table("collection_papers")
        .select("collection_id")
        .eq("paper_id", source_paper_id)
        .execute()
    )

    for row in collection_result.data or []:
        try:
            (
                supabase.table("collection_papers")
                .insert(
                    {
                        "paper_id": target_paper_id,
                        "collection_id": row["collection_id"],
                    }
                )
                .execute()
            )
        except APIError as error:
            if not is_duplicate_key_error(error):
                raise


def copy_item_tags(supabase, source_item_id, target_item_id):
    tag_result = (
        supabase.table("item_tags")
        .select("tag_id")
        .eq("item_id", source_item_id)
        .execute()
    )

    for row in tag_result.data or []:
        try:
            (
                supabase.table("item_tags")
                .insert({"item_id": target_item_id, "tag_id": row["tag_id"]})
                .execute()
            )
        except APIError as error:
            if not is_duplicate_key_error(error):
                raise


def copy_item_collections(supabase, source_item_id, target_item_id):
    collection_result = (
        supabase.table("collection_items")
        .select("collection_id")
        .eq("item_id", source_item_id)
        .execute()
    )

    for row in collection_result.data or []:
        try:
            (
                supabase.table("collection_items")
                .insert(
                    {
                        "item_id": target_item_id,
                        "collection_id": row["collection_id"],
                    }
                )
                .execute()
            )
        except APIError as error:
            if not is_duplicate_key_error(error):
                raise


def ensure_user_owns_item(supabase, user_id, item_id):
    result = (
        supabase.table("items")
        .select("id")
        .eq("id", item_id)
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    if not result.data:
        raise RuntimeError("Item ownership could not be confirmed.")


def fetch_item_storage_paths(supabase, user_id, item_id):
    result = (
        supabase.table("attachments")
        .select("storage_path")
        .eq("item_id", item_id)
        .eq("user_id", user_id)
        .execute()
    )
    return [
        row["storage_path"]
        for row in (result.data or [])
        if is_storage_path(row.get("storage_path"))
    ]


def transfer_item_attachments(
    supabase,
    user_id,
    source_item_id,
    target_item_id,
    keeper,
    duplicate,
):
    conflicts = []
    transferred_fields = {}

    for _, field, label in (
        ("pdf", "pdf_path", "PDF"),
        ("supporting", "supporting_path", "補足資料"),
    ):
        keeper_value = keeper.get(field)
        duplicate_value = duplicate.get(field)
        if keeper_value and duplicate_value and keeper_value != duplicate_value:
            conflicts.append(label)

    if conflicts:
        return transferred_fields, conflicts

    for kind, field, label in (
        ("pdf", "pdf_path", "PDF"),
        ("supporting", "supporting_path", "補足資料"),
    ):
        keeper_value = keeper.get(field)
        duplicate_value = duplicate.get(field)
        if not keeper_value and duplicate_value:
            (
                supabase.table("attachments")
                .update({"item_id": target_item_id})
                .eq("item_id", source_item_id)
                .eq("user_id", user_id)
                .eq("kind", kind)
                .execute()
            )
            transferred_fields[field] = duplicate_value

    return transferred_fields, conflicts


def build_paper_merge_update(keeper, duplicate, preferred_fields=None):
    update_fields = {}
    conflicts = []
    preferred_fields = preferred_fields or {}

    for field in ("title", "authors", "journal", "year", "doi", "url", "status"):
        if preferred_fields.get(field) == "duplicate" and duplicate.get(field):
            update_fields[field] = duplicate[field]
        elif preferred_fields.get(field) == "keeper":
            continue
        elif not keeper.get(field) and duplicate.get(field):
            update_fields[field] = duplicate[field]

    for field, label in (("pdf_path", "PDF"), ("supporting_path", "補足資料")):
        keeper_value = keeper.get(field)
        duplicate_value = duplicate.get(field)
        if preferred_fields.get(field) == "duplicate" and duplicate_value:
            update_fields[field] = duplicate_value
        elif preferred_fields.get(field) == "keeper":
            continue
        elif keeper_value and duplicate_value and keeper_value != duplicate_value:
            conflicts.append(label)
        elif not keeper_value and duplicate_value:
            update_fields[field] = duplicate_value

    keeper_notes = (keeper.get("notes") or "").strip()
    duplicate_notes = (duplicate.get("notes") or "").strip()
    if duplicate_notes and duplicate_notes not in keeper_notes:
        if keeper_notes:
            update_fields["notes"] = f"{keeper_notes}\n\n--- 統合元メモ ---\n{duplicate_notes}"
        else:
            update_fields["notes"] = duplicate_notes

    return update_fields, conflicts


def build_item_merge_update(keeper, duplicate, preferred_fields=None):
    update_fields = {}
    preferred_fields = preferred_fields or {}

    field_map = {
        "title": "title",
        "journal": "publication_title",
        "year": "year",
        "doi": "doi",
        "url": "url",
        "volume": "volume",
        "issue": "issue",
        "pages": "pages",
        "publisher": "publisher",
        "item_type": "item_type",
        "notes": "abstract_note",
    }

    for view_field, item_field in field_map.items():
        if view_field == "notes":
            continue
        if preferred_fields.get(view_field) == "duplicate" and duplicate.get(view_field):
            update_fields[item_field] = duplicate[view_field]
        elif preferred_fields.get(view_field) == "keeper":
            continue
        elif not keeper.get(view_field) and duplicate.get(view_field):
            update_fields[item_field] = duplicate[view_field]

    keeper_notes = (keeper.get("notes") or "").strip()
    duplicate_notes = (duplicate.get("notes") or "").strip()
    if duplicate_notes and duplicate_notes not in keeper_notes:
        if keeper_notes:
            update_fields["abstract_note"] = (
                f"{keeper_notes}\n\n--- 統合元メモ ---\n{duplicate_notes}"
            )
        else:
            update_fields["abstract_note"] = duplicate_notes

    return update_fields


def item_update_to_view_fields(update_fields):
    field_map = {
        "publication_title": "journal",
        "abstract_note": "notes",
    }
    return {
        field_map.get(field, field): value
        for field, value in update_fields.items()
    }


def is_item_backed_paper(row):
    return bool(row.get("item_id"))


def get_paper_reference_ids(row):
    values = [row.get("id"), row.get("item_id")]
    return [value for value in values if value is not None]


def normalize_snapshot_value(value):
    value = normalize_json_value(value)
    if isinstance(value, dict):
        return {
            str(key): normalize_snapshot_value(child_value)
            for key, child_value in value.items()
        }
    if isinstance(value, (list, tuple)):
        return [normalize_snapshot_value(child_value) for child_value in value]
    return value


def create_duplicate_merge_backup(
    supabase,
    user_id,
    keeper,
    duplicate,
    merge_group_id=None,
):
    merge_group_id = merge_group_id or str(uuid.uuid4())
    payload = {
        "user_id": user_id,
        "merge_group_id": merge_group_id,
        "keeper_paper_id": str(keeper.get("id")) if keeper.get("id") is not None else None,
        "duplicate_paper_id": (
            str(duplicate.get("id")) if duplicate.get("id") is not None else None
        ),
        "keeper_item_id": keeper.get("item_id") or None,
        "duplicate_item_id": duplicate.get("item_id") or None,
        "keeper_snapshot": normalize_snapshot_value(keeper),
        "duplicate_snapshot": normalize_snapshot_value(duplicate),
    }
    result = supabase.table("duplicate_merge_backups").insert(payload).execute()
    backup_id = None
    if result.data:
        backup_id = result.data[0].get("id")
    return {"backup_id": backup_id, "merge_group_id": merge_group_id}


def fetch_duplicate_merge_backups(supabase, user_id, limit=50):
    result = (
        supabase.table("duplicate_merge_backups")
        .select(
            "id, merge_group_id, keeper_paper_id, duplicate_paper_id, "
            "keeper_item_id, duplicate_item_id, keeper_snapshot, duplicate_snapshot, created_at"
        )
        .eq("user_id", user_id)
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
    )
    return result.data or []


def is_missing_pdf_annotations_error(error):
    error_text = str(error).lower()
    return (
        "pdf_annotations" in error_text
        and (
            "could not find the table" in error_text
            or "relation" in error_text
            or "does not exist" in error_text
            or "schema cache" in error_text
        )
    )


def fetch_pdf_annotations(supabase, user_id, paper_id):
    if not paper_id:
        return []
    try:
        result = (
            supabase.table("pdf_annotations")
            .select("*")
            .eq("user_id", user_id)
            .eq("paper_id", str(paper_id))
            .order("page_number")
            .order("created_at")
            .execute()
        )
        return result.data or []
    except APIError as error:
        if is_missing_pdf_annotations_error(error):
            return []
        raise


def normalize_annotation_rect(rect):
    if not isinstance(rect, dict):
        return {}

    normalized = {}
    for source_key, target_key in (
        ("x", "rect_x"),
        ("y", "rect_y"),
        ("width", "rect_width"),
        ("height", "rect_height"),
    ):
        try:
            value = float(rect.get(source_key))
        except (TypeError, ValueError):
            return {}
        normalized[target_key] = value

    if (
        normalized["rect_x"] < 0
        or normalized["rect_y"] < 0
        or normalized["rect_width"] <= 0
        or normalized["rect_height"] <= 0
        or normalized["rect_x"] + normalized["rect_width"] > 1
        or normalized["rect_y"] + normalized["rect_height"] > 1
    ):
        return {}
    return normalized


def create_pdf_annotation(
    supabase,
    user_id,
    paper_id,
    page_number,
    annotation_type,
    selected_text="",
    note="",
    color="#fff6db",
    rect=None,
):
    annotation_type = annotation_type if annotation_type in PDF_ANNOTATION_TYPES else "page_note"
    payload = {
        "user_id": user_id,
        "paper_id": str(paper_id),
        "page_number": max(int(page_number or 1), 1),
        "annotation_type": annotation_type,
        "selected_text": normalize_text_db_value(selected_text),
        "note": normalize_text_db_value(note),
        "color": color or "#fff6db",
    }
    payload.update(normalize_annotation_rect(rect))
    return supabase.table("pdf_annotations").insert(payload).execute()


def update_pdf_annotation(
    supabase,
    user_id,
    annotation_id,
    annotation_type,
    selected_text="",
    note="",
    color="#fff6db",
    rect=None,
):
    annotation_type = annotation_type if annotation_type in PDF_ANNOTATION_TYPES else "page_note"
    payload = {
        "annotation_type": annotation_type,
        "selected_text": normalize_text_db_value(selected_text),
        "note": normalize_text_db_value(note),
        "color": color or "#fff6db",
    }
    rect_fields = normalize_annotation_rect(rect)
    if rect_fields:
        payload.update(rect_fields)
    return (
        supabase.table("pdf_annotations")
        .update(payload)
        .eq("id", annotation_id)
        .eq("user_id", user_id)
        .execute()
    )


def delete_pdf_annotation(supabase, user_id, annotation_id):
    return (
        supabase.table("pdf_annotations")
        .delete()
        .eq("id", annotation_id)
        .eq("user_id", user_id)
        .execute()
    )


def restore_keeper_from_merge_backup(supabase, user_id, backup):
    snapshot = backup.get("keeper_snapshot") or {}
    item_id = backup.get("keeper_item_id")
    paper_id = backup.get("keeper_paper_id")

    if item_id:
        ensure_user_owns_item(supabase, user_id, item_id)
        update_fields = build_item_merge_update(snapshot, {})
        for view_field, item_field in {
            "title": "title",
            "journal": "publication_title",
            "year": "year",
            "doi": "doi",
            "url": "url",
            "volume": "volume",
            "issue": "issue",
            "pages": "pages",
            "publisher": "publisher",
            "item_type": "item_type",
            "notes": "abstract_note",
        }.items():
            update_fields[item_field] = snapshot.get(view_field)
        (
            supabase.table("items")
            .update(update_fields)
            .eq("id", item_id)
            .eq("user_id", user_id)
            .execute()
        )
        return {"restored_table": "items", "restored_id": item_id}

    if not paper_id:
        raise ValueError("復元対象の文献IDがバックアップにありません。")

    update_fields = {
        field: snapshot.get(field)
        for field in (
            "title",
            "authors",
            "journal",
            "year",
            "doi",
            "url",
            "status",
            "notes",
            "pdf_path",
            "supporting_path",
            "volume",
            "issue",
            "pages",
            "publisher",
            "item_type",
        )
    }
    (
        supabase.table("papers")
        .update(update_fields)
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )
    return {"restored_table": "papers", "restored_id": paper_id}


def restore_duplicate_from_merge_backup(supabase, user_id, backup):
    snapshot = backup.get("duplicate_snapshot") or {}
    if not snapshot:
        raise ValueError("復元できる統合元スナップショットがありません。")

    next_order = get_next_display_order(supabase, user_id)
    created = create_user_paper(
        supabase,
        user_id,
        snapshot.get("title") or "",
        snapshot.get("authors") or "",
        snapshot.get("journal") or "",
        snapshot.get("year") or 0,
        normalize_doi(snapshot.get("doi")) or None,
        snapshot.get("url") or None,
        snapshot.get("pdf_path") or None,
        snapshot.get("supporting_path") or None,
        snapshot.get("status") or "未読",
        snapshot.get("notes") or "",
        next_order,
        snapshot.get("volume") or "",
        snapshot.get("issue") or "",
        snapshot.get("pages") or "",
        snapshot.get("publisher") or "",
        snapshot.get("item_type") or "journalArticle",
    )
    return {"restored_table": "items" if created.get("item_id") else "papers", "restored_id": created["id"]}


def update_document_citation(supabase, user_id, citation_id, rendered_text=None, context_text=None, sort_order=None):
    fields = {}
    if rendered_text is not None:
        fields["rendered_text"] = rendered_text
    if context_text is not None:
        fields["context_text"] = context_text
    if sort_order is not None:
        fields["sort_order"] = int(sort_order)
    if not fields:
        return None
    return (
        supabase.table("document_citations")
        .update(fields)
        .eq("id", citation_id)
        .execute()
    )


def delete_document_citation(supabase, citation_id):
    return supabase.table("document_citations").delete().eq("id", citation_id).execute()


def merge_duplicate_paper(
    supabase,
    user_id,
    keeper,
    duplicate,
    merge_group_id=None,
    preferred_fields=None,
):
    backup = create_duplicate_merge_backup(
        supabase,
        user_id,
        keeper,
        duplicate,
        merge_group_id=merge_group_id,
    )
    if is_item_backed_paper(keeper) or is_item_backed_paper(duplicate):
        if not is_item_backed_paper(keeper) or not is_item_backed_paper(duplicate):
            raise ValueError("items由来とpapers由来の文献は自動統合できません。")
        result = merge_duplicate_item(
            supabase,
            user_id,
            keeper,
            duplicate,
            preferred_fields=preferred_fields,
        )
        result["backup_ids"] = [backup["backup_id"]] if backup["backup_id"] else []
        result["merge_group_id"] = backup["merge_group_id"]
        return result

    update_fields, conflicts = build_paper_merge_update(
        keeper,
        duplicate,
        preferred_fields=preferred_fields,
    )
    if conflicts:
        raise ValueError(
            " / ".join(conflicts)
            + " が両方の文献にあります。先に残す添付を手動で決めてください。"
        )

    if update_fields:
        (
            supabase.table("papers")
            .update(update_fields)
            .eq("id", keeper["id"])
            .eq("user_id", user_id)
            .execute()
        )

    copy_paper_tags(supabase, duplicate["id"], keeper["id"])
    copy_paper_collections(supabase, duplicate["id"], keeper["id"])
    citation_updates = replace_paper_id_in_document_citations(
        supabase,
        user_id,
        duplicate["id"],
        keeper["id"],
    )

    supabase.table("paper_tags").delete().eq("paper_id", duplicate["id"]).execute()
    supabase.table("collection_papers").delete().eq("paper_id", duplicate["id"]).execute()
    (
        supabase.table("papers")
        .delete()
        .eq("id", duplicate["id"])
        .eq("user_id", user_id)
        .execute()
    )
    remaining = (
        supabase.table("papers")
        .select("id")
        .eq("id", duplicate["id"])
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    if remaining.data:
        raise RuntimeError("統合元の文献が削除されませんでした。権限またはRLSを確認してください。")

    return {
        "citation_updates": citation_updates,
        "updated_fields": update_fields,
        "backup_ids": [backup["backup_id"]] if backup["backup_id"] else [],
        "merge_group_id": backup["merge_group_id"],
    }


def merge_duplicate_item(supabase, user_id, keeper, duplicate, preferred_fields=None):
    ensure_user_owns_item(supabase, user_id, keeper["item_id"])
    ensure_user_owns_item(supabase, user_id, duplicate["item_id"])

    update_fields = build_item_merge_update(
        keeper,
        duplicate,
        preferred_fields=preferred_fields,
    )
    transferred_fields, attachment_conflicts = transfer_item_attachments(
        supabase,
        user_id,
        duplicate["item_id"],
        keeper["item_id"],
        keeper,
        duplicate,
    )
    if attachment_conflicts:
        raise ValueError(
            " / ".join(attachment_conflicts)
            + " が両方の文献にあります。先に残す添付を手動で決めてください。"
        )

    if update_fields:
        (
            supabase.table("items")
            .update(update_fields)
            .eq("id", keeper["item_id"])
            .eq("user_id", user_id)
            .execute()
        )

    copy_item_tags(supabase, duplicate["item_id"], keeper["item_id"])
    copy_item_collections(supabase, duplicate["item_id"], keeper["item_id"])
    citation_updates = replace_paper_id_in_document_citations(
        supabase,
        user_id,
        get_paper_reference_ids(duplicate),
        keeper["id"],
    )

    supabase.table("item_tags").delete().eq("item_id", duplicate["item_id"]).execute()
    supabase.table("collection_items").delete().eq("item_id", duplicate["item_id"]).execute()
    (
        supabase.table("items")
        .delete()
        .eq("id", duplicate["item_id"])
        .eq("user_id", user_id)
        .execute()
    )
    remaining = (
        supabase.table("items")
        .select("id")
        .eq("id", duplicate["item_id"])
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    if remaining.data:
        raise RuntimeError("統合元の文献が削除されませんでした。権限またはRLSを確認してください。")

    updated_fields = item_update_to_view_fields(update_fields)
    updated_fields.update(transferred_fields)
    return {"citation_updates": citation_updates, "updated_fields": updated_fields}


def sort_papers_dataframe(df, sort_option, added_oldest_first=False):
    if df.empty:
        return df

    sorted_df = df.copy()

    if sort_option == "追加順" and "display_order" in sorted_df.columns:
        sorted_df = sorted_df.sort_values(
            by="display_order",
            ascending=added_oldest_first,
        )
    elif sort_option == "年（新しい順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=False)
    elif sort_option == "年（古い順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=True)
    elif sort_option == "タイトル":
        sorted_df = sorted_df.sort_values(by="title", ascending=True)
    elif sort_option == "ステータス":
        sorted_df = sorted_df.sort_values(by="status", ascending=True)

    sorted_df = sorted_df.reset_index(drop=True)
    sorted_df["ref_no"] = sorted_df.index + 1
    return sorted_df


def get_csl_style_path(style):
    if citeproc_styles is None:
        return None
    style_id = CSL_STYLE_OPTIONS.get(style, style)
    style_id = str(style_id or "").strip().lower()
    if not style_id:
        return None
    if style_id.endswith(".csl"):
        style_id = style_id[:-4]
    style_path = os.path.join(
        os.path.dirname(citeproc_styles.__file__),
        "styles",
        f"{style_id}.csl",
    )
    return style_path if os.path.exists(style_path) else None


def list_available_csl_styles(query="", limit=80):
    styles = []
    if citeproc_styles is None:
        return styles

    styles_dir = os.path.join(os.path.dirname(citeproc_styles.__file__), "styles")
    if not os.path.isdir(styles_dir):
        return styles

    normalized_query = str(query or "").strip().lower()
    try:
        names = os.listdir(styles_dir)
    except OSError:
        return styles

    for name in names:
        if not name.endswith(".csl"):
            continue
        style_id = name[:-4]
        if normalized_query and normalized_query not in style_id.lower():
            continue
        styles.append(style_id)

    preferred = list(CSL_STYLE_OPTIONS.values())
    styles = sorted(set(styles), key=lambda value: (value not in preferred, value))
    return styles[: max(int(limit or 80), 1)]


def parse_csl_authors(authors):
    csl_authors = []
    for name in normalize_bibtex_authors(authors):
        text = re.sub(r"\s+", " ", name or "").strip()
        if not text:
            continue
        if "," in text:
            family, given = [part.strip() for part in text.split(",", maxsplit=1)]
        else:
            parts = text.split()
            family = parts[-1]
            given = " ".join(parts[:-1])
        author = {"family": family}
        if given:
            author["given"] = given
        csl_authors.append(author)
    return csl_authors


def paper_to_csl_json(row):
    year = row.get("year")
    try:
        year = int(year) if year not in (None, "") else None
    except (TypeError, ValueError):
        year = None
    item_type = row.get("item_type") or ("journalArticle" if row.get("journal") else "webpage")
    csl_type = {
        "journalArticle": "article-journal",
        "book": "book",
        "bookSection": "chapter",
        "webpage": "webpage",
        "thesis": "thesis",
        "report": "report",
    }.get(item_type, "article-journal")
    item_id = str(row.get("item_id") or row.get("id") or "item")
    data = {
        "id": item_id,
        "type": csl_type,
        "title": row.get("title") or "",
        "container-title": row.get("journal") or "",
        "volume": row.get("volume") or "",
        "issue": row.get("issue") or "",
        "page": row.get("pages") or "",
        "publisher": row.get("publisher") or "",
        "DOI": normalize_doi(row.get("doi")) or "",
        "URL": row.get("url") or "",
        "author": parse_csl_authors(row.get("authors")),
    }
    if year:
        data["issued"] = {"date-parts": [[year]]}
    return {key: value for key, value in data.items() if value not in ("", [], None)}


def format_csl_bibliography_entry(row, style="APA"):
    if not all([Citation, CitationItem, CitationStylesBibliography, CitationStylesStyle, CiteProcJSON, plain]):
        return None
    style_path = get_csl_style_path(style)
    if not style_path:
        return None
    try:
        csl_item = paper_to_csl_json(row)
        source = CiteProcJSON([csl_item])
        csl_style = CitationStylesStyle(style_path, validate=False)
        bibliography = CitationStylesBibliography(csl_style, source, plain)
        citation = Citation([CitationItem(csl_item["id"])])
        bibliography.register(citation)
        entries = bibliography.bibliography()
    except Exception:
        return None
    if not entries:
        return None
    return append_missing_csl_doi(str(entries[0]).strip(), row, style)


def append_missing_csl_doi(text, row, style="APA"):
    doi = normalize_doi(row.get("doi"))
    if not doi or doi.lower() in text.lower():
        return text
    if str(style or "").lower() in {"apa", "elsevier-harvard", "chicago-author-date"}:
        return f"{text} https://doi.org/{doi}"
    return f"{text} doi: {doi}"


def make_word_citation_fallback(row, style="APA"):
    authors = normalize_author_text(row.get("authors", ""))
    year = row.get("year", "")
    title = row.get("title", "")
    journal = row.get("journal", "")
    doi = normalize_doi(row.get("doi", ""))
    volume = row.get("volume", "")
    issue = row.get("issue", "")
    pages = row.get("pages", "")
    publication = journal
    if volume:
        publication += f", {volume}"
        if issue:
            publication += f"({issue})"
    if pages:
        publication += f", {pages}"

    if style == "APA":
        citation = f"{authors} ({year}). {title}. {publication}."
        if doi:
            citation += f" https://doi.org/{doi}"
    elif style == "Vancouver":
        citation = f"{authors}. {title}. {publication}. {year}."
        if doi:
            citation += f" doi:{doi}"
    elif style == "Nature":
        citation = f"{authors} {title}. {publication} ({year})."
        if doi:
            citation += f" https://doi.org/{doi}"
    else:
        citation = f"{authors} ({year}). {title}. {publication}."

    return citation


def make_word_citation(row, style="APA"):
    csl_entry = format_csl_bibliography_entry(row, style=style)
    if csl_entry:
        return csl_entry
    return make_word_citation_fallback(row, style=style)


def make_bibtex_key(row):
    authors = row.get("authors") or "unknown"
    first_author = re.split(r",| and ", authors, maxsplit=1)[0].strip()
    author_key = re.sub(r"[^A-Za-z0-9]+", "", first_author) or "unknown"
    year_key = re.sub(r"[^0-9]+", "", str(row.get("year") or "")) or "nodate"
    title_words = re.findall(r"[A-Za-z0-9]+", row.get("title") or "")
    title_key = title_words[0] if title_words else "untitled"
    return f"{author_key}{year_key}{title_key}"


def normalize_bibtex_authors(authors):
    names = [name.strip() for name in re.split(r"\s+and\s+|;", authors or "") if name.strip()]
    if len(names) == 1:
        names = [name.strip() for name in (authors or "").split(",") if name.strip()]
    return names


def normalize_author_text(authors):
    return ", ".join(normalize_bibtex_authors(authors))


def normalize_bibtex_doi(doi):
    return normalize_doi(doi)


def escape_bibtex_value(value):
    text = str(value or "")
    return text.replace("\\", "\\textbackslash{}").replace("{", "\\{").replace("}", "\\}")


def make_bibtex_entry(row):
    entry_type = "article" if row.get("journal") else "misc"
    fields = [
        ("title", row.get("title")),
        ("author", " and ".join(normalize_bibtex_authors(row.get("authors")))),
        ("journal", row.get("journal")),
        ("year", row.get("year")),
        ("volume", row.get("volume")),
        ("number", row.get("issue")),
        ("pages", row.get("pages")),
        ("publisher", row.get("publisher")),
        ("doi", normalize_bibtex_doi(row.get("doi"))),
        ("url", row.get("url")),
    ]
    lines = [f"@{entry_type}{{{make_bibtex_key(row)},"]
    for field, value in fields:
        if value not in (None, ""):
            lines.append(f"  {field} = {{{escape_bibtex_value(value)}}},")
    if len(lines) > 1:
        lines[-1] = lines[-1].rstrip(",")
    lines.append("}")
    return "\n".join(lines)


def make_ris_entry(row):
    ris_type = "JOUR" if row.get("journal") else "GEN"
    lines = [f"TY  - {ris_type}"]
    for author in normalize_bibtex_authors(row.get("authors")):
        lines.append(f"AU  - {author}")
    field_map = [
        ("TI", row.get("title")),
        ("T2", row.get("journal")),
        ("PY", row.get("year")),
        ("VL", row.get("volume")),
        ("IS", row.get("issue")),
        ("SP", row.get("pages")),
        ("PB", row.get("publisher")),
        ("DO", normalize_bibtex_doi(row.get("doi"))),
        ("UR", row.get("url")),
    ]
    for tag, value in field_map:
        if value not in (None, ""):
            lines.append(f"{tag}  - {value}")
    lines.append("ER  -")
    return "\n".join(lines)


def export_to_bibtex_text(papers):
    return "\n\n".join(make_bibtex_entry(paper) for paper in papers or [])


def export_to_ris_text(papers):
    return "\n\n".join(make_ris_entry(paper) for paper in papers or [])


def parse_bibtex_entries(text):
    entries = []
    for match in re.finditer(r"@\w+\s*\{\s*[^,]+,(.*?)(?=\n\s*@|\Z)", text or "", re.DOTALL):
        body = match.group(1)
        fields = {}
        for field_match in re.finditer(
            r"(\w+)\s*=\s*(?:\{(.*?)\}|\"(.*?)\")\s*,?",
            body,
            re.DOTALL,
        ):
            name = field_match.group(1).lower()
            value = field_match.group(2) if field_match.group(2) is not None else field_match.group(3)
            fields[name] = re.sub(r"\s+", " ", value or "").strip()
        if fields:
            entries.append(
                {
                    "title": fields.get("title", ""),
                    "authors": ", ".join(normalize_bibtex_authors(fields.get("author", ""))),
                    "journal": fields.get("journal") or fields.get("booktitle", ""),
                    "year": fields.get("year") or 0,
                    "doi": normalize_bibtex_doi(fields.get("doi", "")),
                    "url": fields.get("url", ""),
                    "volume": fields.get("volume", ""),
                    "issue": fields.get("number", ""),
                    "pages": fields.get("pages", ""),
                    "publisher": fields.get("publisher", ""),
                }
            )
    return entries


def parse_ris_entries(text):
    entries = []
    current = {}
    authors = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        line_match = re.match(r"^([A-Z0-9]{2})\s+-\s*(.*)$", line, flags=re.IGNORECASE)
        if not line_match:
            continue
        tag = line_match.group(1).strip().upper()
        value = line_match.group(2).strip()
        if tag == "TY":
            current = {}
            authors = []
        elif tag == "AU":
            authors.append(value)
        elif tag in {"TI", "T1"}:
            current["title"] = value
        elif tag in {"T2", "JO", "JF"}:
            current["journal"] = value
        elif tag in {"PY", "Y1"}:
            year_match = re.search(r"\d{4}", value)
            current["year"] = year_match.group(0) if year_match else value
        elif tag == "VL":
            current["volume"] = value
        elif tag == "IS":
            current["issue"] = value
        elif tag in {"SP", "EP"}:
            current["pages"] = (
                f"{current.get('pages', '')}-{value}".strip("-")
                if tag == "EP"
                else value
            )
        elif tag == "PB":
            current["publisher"] = value
        elif tag == "DO":
            current["doi"] = normalize_bibtex_doi(value)
        elif tag == "UR":
            current["url"] = value
        elif tag == "ER":
            current["authors"] = ", ".join(authors)
            entries.append(
                {
                    "title": current.get("title", ""),
                    "authors": current.get("authors", ""),
                    "journal": current.get("journal", ""),
                    "year": current.get("year") or 0,
                    "doi": current.get("doi", ""),
                    "url": current.get("url", ""),
                    "volume": current.get("volume", ""),
                    "issue": current.get("issue", ""),
                    "pages": current.get("pages", ""),
                    "publisher": current.get("publisher", ""),
                }
            )
            current = {}
            authors = []
    return [entry for entry in entries if entry.get("title") or entry.get("doi")]


def extract_doi_from_pdf_bytes(pdf_bytes):
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text_parts = []
        for page in reader.pages[:3]:
            text_parts.append(page.extract_text() or "")
        return extract_doi_from_text(" ".join(text_parts))
    except Exception:
        return ""


def extract_title_from_pdf_bytes(pdf_bytes):
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        metadata_title = (reader.metadata or {}).get("/Title") or ""
        metadata_title = re.sub(r"\s+", " ", str(metadata_title)).strip()
        if metadata_title and len(metadata_title) >= 8:
            return metadata_title

        first_page_text = reader.pages[0].extract_text() if reader.pages else ""
    except Exception:
        return ""

    candidates = []
    for raw_line in (first_page_text or "").splitlines()[:30]:
        line = re.sub(r"\s+", " ", raw_line).strip()
        if len(line) < 12 or DOI_RE.search(line):
            continue
        if re.search(r"^(abstract|keywords|introduction|references)\b", line, re.I):
            continue
        if re.search(r"^(journal|volume|copyright|received|accepted)\b", line, re.I):
            continue
        candidates.append(line)
    if not candidates:
        return ""
    return max(candidates[:8], key=len)[:300]


def extract_text_from_pdf_bytes(pdf_bytes, max_pages=20, max_chars=PDF_SUMMARY_MAX_CHARS):
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception:
        return ""
    text_parts = []
    for page in reader.pages[:max_pages]:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            continue
        if sum(len(part) for part in text_parts) >= max_chars:
            break
    text = "\n".join(text_parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:max_chars].strip()


def extract_pdf_summary_sections(pdf_text):
    text = (pdf_text or "").strip()
    if not text:
        return {"abstract": "", "introduction": "", "conclusion": "", "fallback": ""}

    heading_patterns = {
        "abstract": r"(?:^|\n)\s*(?:abstract|summary)\s*(?:\n|$)",
        "introduction": r"(?:^|\n)\s*(?:\d+\.?\s*)?introduction\s*(?:\n|$)",
        "conclusion": (
            r"(?:^|\n)\s*(?:\d+\.?\s*)?"
            r"(?:conclusion|conclusions|summary and conclusions|concluding remarks)\s*(?:\n|$)"
        ),
    }
    stop_pattern = re.compile(
        r"(?:^|\n)\s*(?:\d+\.?\s*)?"
        r"(?:abstract|keywords?|introduction|methods?|materials and methods|results|discussion|"
        r"conclusion|conclusions|summary|acknowledg(?:e)?ments?|references|bibliography)\s*(?:\n|$)",
        re.IGNORECASE,
    )

    def section_after(pattern, limit):
        match = re.search(pattern, text, re.IGNORECASE)
        if not match:
            return ""
        start = match.end()
        next_match = stop_pattern.search(text, start)
        end = next_match.start() if next_match else min(len(text), start + limit)
        section = text[start:end].strip()
        section = re.sub(r"\n+", "\n", section)
        return section[:limit].strip()

    return {
        "abstract": section_after(heading_patterns["abstract"], 8000),
        "introduction": section_after(heading_patterns["introduction"], 16000),
        "conclusion": section_after(heading_patterns["conclusion"], 12000),
        "fallback": text[:24000],
    }


def summarize_paper_with_gemini(
    api_key,
    paper,
    sections,
    model="gemini-2.5-flash",
):
    if genai is None:
        raise RuntimeError("google-genai がインストールされていません。requirements.txt を確認してください。")
    if not api_key:
        raise ValueError("GEMINI_API_KEY が設定されていません。")

    source_sections = {
        "Abstract": sections.get("abstract") or "",
        "Introduction": sections.get("introduction") or "",
        "Conclusion/Summary": sections.get("conclusion") or "",
    }
    if not any(value.strip() for value in source_sections.values()):
        source_sections["本文冒頭"] = sections.get("fallback") or ""
    source_text = "\n\n".join(
        f"## {label}\n{value.strip()}"
        for label, value in source_sections.items()
        if value.strip()
    )
    if not source_text.strip():
        raise ValueError("PDFから要約に使える本文を抽出できませんでした。")

    prompt = f"""
あなたは研究論文を読む日本語の研究支援アシスタントです。
以下の論文情報とPDF抽出テキストをもとに、本文に書かれている範囲だけで要約してください。
推測で補わず、不明な点は「不明」と書いてください。

論文タイトル: {paper.get("title") or "不明"}
著者: {paper.get("authors") or "不明"}
掲載誌/年: {paper.get("journal") or "不明"} / {paper.get("year") or "不明"}
DOI: {paper.get("doi") or "不明"}

出力形式:
### 全体要約
- 3〜5点

### Abstractの要点
- 2〜4点

### Introductionの要点
- 2〜4点

### まとめ・結論の要点
- 2〜4点

### 読むときの注目点
- 2〜4点

PDF抽出テキスト:
{source_text[:PDF_SUMMARY_MAX_CHARS]}
""".strip()

    client = genai.Client(api_key=api_key)
    config = None
    if genai_types is not None:
        config = genai_types.GenerateContentConfig(temperature=0.2)
    response = client.models.generate_content(
        model=model or "gemini-2.5-flash",
        contents=prompt,
        config=config,
    )
    return (getattr(response, "text", None) or "").strip()


def export_to_word_bytes(papers):
    doc = Document()
    doc.add_heading("参考文献", 0)

    for index, paper in enumerate(papers, start=1):
        text = (
            f"[{index}] {paper.get('authors', '')} ({paper.get('year', '')}). "
            f"{paper.get('title', '')}. {paper.get('journal', '')}."
        )
        doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def upload_file_to_storage(
    supabase,
    uploaded_file,
    user_id,
    folder,
    default_ext=".bin",
    default_content_type="application/octet-stream",
):
    safe_name = make_safe_storage_filename(
        getattr(uploaded_file, "name", "attachment"),
        default_ext=default_ext,
    )
    storage_path = f"{user_id}/{folder}/{safe_name}"
    content_type = getattr(uploaded_file, "type", None) or default_content_type

    supabase.storage.from_(BUCKET_NAME).upload(
        path=storage_path,
        file=uploaded_file.read(),
        file_options={"content-type": content_type},
    )
    return storage_path


def upload_pdf_to_storage(supabase, pdf_file, user_id):
    return upload_file_to_storage(
        supabase,
        pdf_file,
        user_id,
        "pdfs",
        default_ext=".pdf",
        default_content_type="application/pdf",
    )


def upload_supporting_file_to_storage(supabase, supporting_file, user_id):
    return upload_file_to_storage(supabase, supporting_file, user_id, "supporting")


def create_pdf_signed_url(supabase, storage_path, expires_in=3600):
    if not isinstance(storage_path, str) or not storage_path.strip():
        return None

    response = (
        supabase.storage.from_(BUCKET_NAME).create_signed_url(
            storage_path,
            expires_in,
        )
    )

    if isinstance(response, dict):
        return response.get("signedURL") or response.get("signedUrl")
    return None


def delete_pdf_from_storage(supabase, storage_path):
    if is_storage_path(storage_path):
        supabase.storage.from_(BUCKET_NAME).remove([storage_path])


def download_pdf_from_storage(supabase, storage_path):
    if not has_attachment_path(storage_path):
        return None

    response = supabase.storage.from_(BUCKET_NAME).download(storage_path)
    if isinstance(response, bytes):
        return response
    if isinstance(response, bytearray):
        return bytes(response)
    if hasattr(response, "content"):
        return response.content
    if hasattr(response, "read"):
        return response.read()
    return bytes(response)


def make_pdf_archive_filename(paper, index, used_names=None):
    used_names = used_names if used_names is not None else set()
    title = normalize_text_db_value((paper or {}).get("title")).strip()
    fallback = normalize_text_db_value((paper or {}).get("id")).strip() or "paper"
    base = title or fallback
    base = unicodedata.normalize("NFKC", base)
    base = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", base)
    base = re.sub(r"\s+", " ", base).strip(" ._")
    base = base[:120].strip(" ._") or "paper"

    candidate = f"{index:03d}_{base}.pdf"
    counter = 2
    while candidate in used_names:
        candidate = f"{index:03d}_{base}_{counter}.pdf"
        counter += 1
    used_names.add(candidate)
    return candidate


def build_pdf_download_zip(supabase, papers):
    buffer = io.BytesIO()
    downloaded = []
    failed = []
    used_names = set()

    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, paper in enumerate(papers or [], start=1):
            storage_path = (paper or {}).get("pdf_path")
            if not has_attachment_path(storage_path):
                continue
            try:
                pdf_bytes = download_pdf_from_storage(supabase, storage_path)
            except Exception as error:
                failed.append(
                    {
                        "title": (paper or {}).get("title") or "Untitled",
                        "error": str(error),
                    }
                )
                continue
            if not pdf_bytes:
                failed.append(
                    {
                        "title": (paper or {}).get("title") or "Untitled",
                        "error": "empty_pdf",
                    }
                )
                continue

            filename = make_pdf_archive_filename(paper, index, used_names)
            archive.writestr(filename, pdf_bytes)
            downloaded.append(filename)

    return {
        "bytes": buffer.getvalue(),
        "count": len(downloaded),
        "filenames": downloaded,
        "failed": failed,
    }


def get_or_create_tag_id(supabase, user_id, tag_name):
    tag_result = (
        supabase.table("tags")
        .select("id")
        .eq("name", tag_name)
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )

    if tag_result.data:
        return tag_result.data[0]["id"]

    new_tag = supabase.table("tags").insert({"name": tag_name, "user_id": user_id}).execute()
    return new_tag.data[0]["id"]


def save_tags_for_paper(supabase, user_id, paper_id, tags_text):
    for tag_name in normalize_tag_input(tags_text):
        tag_id = get_or_create_tag_id(supabase, user_id, tag_name)
        supabase.table("paper_tags").upsert(
            {"paper_id": paper_id, "tag_id": tag_id}
        ).execute()


def save_tags_for_item(supabase, user_id, item_id, tags_text):
    if not item_id:
        return
    for tag_name in normalize_tag_input(tags_text):
        tag_id = get_or_create_tag_id(supabase, user_id, tag_name)
        supabase.table("item_tags").upsert(
            {"item_id": item_id, "tag_id": str(tag_id)}
        ).execute()


def replace_tags_for_paper(supabase, user_id, paper_id, item_id, tags_text):
    legacy_paper_id = paper_id if paper_id and not normalize_uuid_text(paper_id) else None
    if legacy_paper_id:
        supabase.table("paper_tags").delete().eq("paper_id", legacy_paper_id).execute()

    if item_id:
        try:
            supabase.table("item_tags").delete().eq("item_id", item_id).execute()
        except APIError as error:
            if not is_missing_relation_error(error):
                raise
        save_tags_for_item(supabase, user_id, item_id, tags_text)
    elif paper_id:
        save_tags_for_paper(supabase, user_id, paper_id, tags_text)


def get_tag_map_for_papers(supabase, papers_or_ids):
    if not papers_or_ids:
        return {}

    if isinstance(papers_or_ids[0], dict):
        papers = papers_or_ids
        paper_ids = [
            normalized_id
            for normalized_id in (normalize_optional_id(row.get("id")) for row in papers)
            if normalized_id
        ]
        item_ids = [
            normalized_id
            for normalized_id in (
                normalize_optional_id(row.get("item_id")) for row in papers
            )
            if normalized_id
        ]
    else:
        paper_ids = [
            normalized_id
            for normalized_id in (normalize_optional_id(paper_id) for paper_id in papers_or_ids)
            if normalized_id
        ]
        item_ids = []

    paper_tag_result = None
    if paper_ids:
        try:
            paper_tag_result = (
                supabase.table("paper_tags")
                .select("paper_id, tag_id")
                .in_("paper_id", paper_ids)
                .execute()
            )
        except APIError:
            return {}
    item_tag_result = None
    if item_ids:
        try:
            item_tag_result = (
                supabase.table("item_tags")
                .select("item_id, tag_id")
                .in_("item_id", item_ids)
                .execute()
            )
        except APIError:
            return {}

    paper_tags = paper_tag_result.data if paper_tag_result else []
    item_tags = item_tag_result.data if item_tag_result else []
    if not paper_tags and not item_tags:
        return {}

    tag_ids = sorted(
        {
            tag_id
            for tag_id in (
                normalize_uuid_text(row.get("tag_id"))
                for row in paper_tags + item_tags
            )
            if tag_id
        }
    )
    if not tag_ids:
        return {}

    try:
        tag_result = supabase.table("tags").select("id, name").in_("id", tag_ids).execute()
    except APIError:
        return {}
    tag_name_map = {str(row["id"]): row["name"] for row in (tag_result.data or [])}

    tag_map = {paper_id: [] for paper_id in paper_ids}
    for row in paper_tags:
        tag_name = tag_name_map.get(str(row["tag_id"]))
        if tag_name:
            tag_map.setdefault(str(row["paper_id"]), []).append(tag_name)
    for row in item_tags:
        tag_name = tag_name_map.get(str(row["tag_id"]))
        if tag_name:
            tag_map.setdefault(str(row["item_id"]), []).append(tag_name)

    return tag_map


def update_item_display_order(supabase, user_id, item_id, display_order):
    item_result = (
        supabase.table("items")
        .select("extra")
        .eq("id", item_id)
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    extra = (item_result.data or [{}])[0].get("extra") or {}
    extra["legacy_display_order"] = str(display_order)
    (
        supabase.table("items")
        .update({"extra": extra})
        .eq("id", item_id)
        .eq("user_id", user_id)
        .execute()
    )


def move_paper(supabase, user_id, paper_id, display_order, direction, item_id=None):
    if item_id:
        result = fetch_user_papers(
            supabase,
            user_id,
            columns="id, item_id, display_order",
        )
        papers = [
            row
            for row in (result.data or [])
            if row.get("display_order") is not None
        ]
        papers.sort(key=lambda row: int(row["display_order"]))
        index = next(
            (
                position
                for position, row in enumerate(papers)
                if str(row.get("item_id")) == str(item_id)
            ),
            None,
        )
        if index is None:
            return
        neighbor_index = index - 1 if direction == "up" else index + 1
        if neighbor_index < 0 or neighbor_index >= len(papers):
            return

        current = papers[index]
        neighbor = papers[neighbor_index]
        update_item_display_order(
            supabase,
            user_id,
            current["item_id"],
            neighbor["display_order"],
        )
        update_item_display_order(
            supabase,
            user_id,
            neighbor["item_id"],
            current["display_order"],
        )
        return

    operator = "lt" if direction == "up" else "gt"
    descending = direction == "up"

    neighbor_result = (
        getattr(
            supabase.table("papers")
            .select("id, display_order")
            .eq("user_id", user_id),
            operator,
        )("display_order", display_order)
        .order("display_order", desc=descending)
        .limit(1)
        .execute()
    )

    if not neighbor_result.data:
        return

    neighbor = neighbor_result.data[0]

    (
        supabase.table("papers")
        .update({"display_order": neighbor["display_order"]})
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )

    (
        supabase.table("papers")
        .update({"display_order": display_order})
        .eq("id", neighbor["id"])
        .eq("user_id", user_id)
        .execute()
    )


def update_paper_details(
    supabase,
    user_id,
    paper_id,
    status,
    notes,
    url=None,
    item_id=None,
    doi=None,
    volume=None,
    issue=None,
    pages=None,
    publisher=None,
):
    doi_provided = doi is not None
    url_provided = url is not None
    status = normalize_text_db_value(status)
    notes = normalize_text_db_value(notes)
    doi = normalize_optional_db_value(doi) if doi_provided else None
    url = normalize_optional_db_value(url) if url_provided else None

    if item_id:
        item_result = (
            supabase.table("items")
            .select("extra")
            .eq("id", item_id)
            .eq("user_id", user_id)
            .limit(1)
            .execute()
        )
        extra = (item_result.data or [{}])[0].get("extra") or {}
        extra["legacy_status"] = status
        fields = {"abstract_note": notes, "extra": extra}
        if doi_provided:
            fields["doi"] = doi
        if url_provided:
            fields["url"] = url
        metadata_fields = {
            "volume": volume,
            "issue": issue,
            "pages": pages,
            "publisher": publisher,
        }
        for field, value in metadata_fields.items():
            if value is not None:
                fields[field] = normalize_optional_db_value(value)
        try:
            (
                supabase.table("items")
                .update(fields)
                .eq("id", item_id)
                .eq("user_id", user_id)
                .execute()
            )
        except APIError as error:
            if not is_missing_metadata_column_error(error):
                raise
            for field in ITEM_METADATA_COLUMNS:
                fields.pop(field, None)
            (
                supabase.table("items")
                .update(fields)
                .eq("id", item_id)
                .eq("user_id", user_id)
                .execute()
            )
        return

    fields = {"status": status, "notes": notes}
    if doi_provided:
        fields["doi"] = doi
    if url_provided:
        fields["url"] = url

    (
        supabase.table("papers")
        .update(fields)
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )


def replace_item_attachment(supabase, user_id, item_id, kind, storage_path):
    if not item_id or storage_path is None:
        return

    ensure_user_owns_item(supabase, user_id, item_id)

    if not is_storage_path(storage_path):
        (
            supabase.table("attachments")
            .delete()
            .eq("item_id", item_id)
            .eq("user_id", user_id)
            .eq("kind", kind)
            .execute()
        )
        return

    existing = (
        supabase.table("attachments")
        .select("id")
        .eq("item_id", item_id)
        .eq("user_id", user_id)
        .eq("kind", kind)
        .limit(1)
        .execute()
    )
    if existing.data:
        (
            supabase.table("attachments")
            .update({"storage_path": storage_path})
            .eq("id", existing.data[0]["id"])
            .eq("user_id", user_id)
            .execute()
        )
    else:
        (
            supabase.table("attachments")
            .insert(
                {
                    "item_id": item_id,
                    "user_id": user_id,
                    "kind": kind,
                    "storage_path": storage_path,
                }
            )
            .execute()
        )


def update_paper_files(
    supabase,
    user_id,
    paper_id,
    pdf_path=None,
    supporting_path=None,
    item_id=None,
):
    if item_id:
        if pdf_path is not None:
            replace_item_attachment(supabase, user_id, item_id, "pdf", pdf_path)
        if supporting_path is not None:
            replace_item_attachment(
                supabase,
                user_id,
                item_id,
                "supporting",
                supporting_path,
            )
        return

    fields = {}
    if pdf_path is not None:
        fields["pdf_path"] = pdf_path
    if supporting_path is not None:
        fields["supporting_path"] = supporting_path

    if not fields:
        return

    (
        supabase.table("papers")
        .update(fields)
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )


def delete_paper(supabase, user_id, row, delete_files=True):
    if is_item_backed_paper(row):
        return delete_item_backed_paper(supabase, user_id, row, delete_files=delete_files)

    storage_errors = []
    pdf_path = row.get("pdf_path")
    supporting_path = row.get("supporting_path")

    supabase.table("paper_tags").delete().eq("paper_id", row["id"]).execute()
    supabase.table("collection_papers").delete().eq("paper_id", row["id"]).execute()
    (
        supabase.table("papers")
        .delete()
        .eq("id", row["id"])
        .eq("user_id", user_id)
        .execute()
    )
    remaining = (
        supabase.table("papers")
        .select("id")
        .eq("id", row["id"])
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    if remaining.data:
        raise RuntimeError("文献が削除されませんでした。権限またはRLSを確認してください。")

    if delete_files:
        for label, storage_path in (("PDF", pdf_path), ("補足資料", supporting_path)):
            if not is_storage_path(storage_path):
                continue
            try:
                delete_pdf_from_storage(supabase, storage_path)
            except Exception as error:
                storage_errors.append(f"{label}: {error}")

    return {"storage_errors": storage_errors}


def delete_item_backed_paper(supabase, user_id, row, delete_files=True):
    storage_errors = []
    ensure_user_owns_item(supabase, user_id, row["item_id"])
    storage_paths = fetch_item_storage_paths(supabase, user_id, row["item_id"])

    supabase.table("item_tags").delete().eq("item_id", row["item_id"]).execute()
    supabase.table("collection_items").delete().eq("item_id", row["item_id"]).execute()
    (
        supabase.table("items")
        .delete()
        .eq("id", row["item_id"])
        .eq("user_id", user_id)
        .execute()
    )
    remaining = (
        supabase.table("items")
        .select("id")
        .eq("id", row["item_id"])
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    if remaining.data:
        raise RuntimeError("文献が削除されませんでした。権限またはRLSを確認してください。")

    if delete_files:
        for storage_path in storage_paths:
            try:
                delete_pdf_from_storage(supabase, storage_path)
            except Exception as error:
                storage_errors.append(f"{storage_path}: {error}")

    return {"storage_errors": storage_errors}
