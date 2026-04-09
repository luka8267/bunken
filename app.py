import io
import hashlib
import hmac
import os
import secrets
import time

import pandas as pd
import requests
import streamlit as st
from docx import Document
from postgrest.exceptions import APIError
from supabase import Client, create_client

BUCKET_NAME = "paper-pdfs"
READING_STATUSES = ["未読", "読書中", "読了", "再読したい", "引用予定"]
SORT_OPTIONS = ["追加順", "年（新しい順）", "年（古い順）", "タイトル", "ステータス"]
DOI_FORM_FIELDS = ("title", "authors", "journal", "year")
ADMIN_RECOVERY_SECRET_NAME = "ADMIN_RECOVERY_TOKEN"


# -----------------------------
# Supabase 接続
# -----------------------------
supabase: Client = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"],
)


# -----------------------------
# 共通処理
# -----------------------------
def get_current_user_id():
    return st.session_state["user_id"]


@st.cache_resource
def get_supabase_admin():
    service_role_key = st.secrets.get("SUPABASE_SERVICE_ROLE_KEY")
    if not service_role_key:
        return None

    return create_client(
        st.secrets["SUPABASE_URL"],
        service_role_key,
    )


def has_service_role_access():
    return get_supabase_admin() is not None


def normalize_doi(doi):
    return (doi or "").strip()


def normalize_username(username):
    return (username or "").strip()


def normalize_tag_input(tags_text):
    seen = set()
    normalized = []
    for tag in (tags_text or "").split(","):
        value = tag.strip()
        if value and value not in seen:
            seen.add(value)
            normalized.append(value)
    return normalized


def hash_password(password):
    salt = secrets.token_hex(16)
    hashed = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        100_000,
    )
    return f"{salt}${hashed.hex()}"


def verify_password(password, password_hash):
    if not password_hash or "$" not in password_hash:
        return False

    salt, stored_hash = password_hash.split("$", maxsplit=1)
    computed_hash = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        100_000,
    ).hex()
    return hmac.compare_digest(computed_hash, stored_hash)


def migrate_legacy_password(user_id, password):
    try:
        (
            supabase.table("users")
            .update({"password_hash": hash_password(password), "password": None})
            .eq("id", user_id)
            .execute()
        )
        return True
    except APIError:
        try:
            (
                supabase.table("users")
                .update({"password_hash": hash_password(password)})
                .eq("id", user_id)
                .execute()
            )
            return True
        except APIError:
            return False


def fetch_user_for_auth(username):
    select_candidates = [
        "id, username, password_hash, password",
        "id, username, password_hash",
        "id, username, password",
    ]

    for select_clause in select_candidates:
        try:
            return (
                supabase.table("users")
                .select(select_clause)
                .eq("username", username)
                .limit(1)
                .execute()
            )
        except APIError:
            continue

    raise APIError({"message": "users table is missing required auth columns"})


def create_user_with_password_hash(username, password):
    insert_payloads = [
        {
            "username": username,
            "password": None,
            "password_hash": hash_password(password),
        },
        {
            "username": username,
            "password_hash": hash_password(password),
        },
    ]

    last_error = None
    for payload in insert_payloads:
        try:
            return supabase.table("users").insert(payload).execute()
        except APIError as error:
            last_error = error

    if last_error:
        raise last_error


def get_admin_recovery_token():
    return st.secrets.get(ADMIN_RECOVERY_SECRET_NAME)


def reset_user_password_by_admin(username, new_password):
    admin_client = get_supabase_admin()
    if admin_client is None:
        raise RuntimeError("service role key is not configured")

    normalized_username = normalize_username(username)
    password_hash = hash_password(new_password)

    return (
        admin_client.table("users")
        .update({"password_hash": password_hash})
        .eq("username", normalized_username)
        .execute()
    )


def authenticate_user(username, password):
    normalized_username = normalize_username(username)
    result = fetch_user_for_auth(normalized_username)

    if not result.data:
        return None, "not_found"

    user = result.data[0]
    password_hash = user.get("password_hash")

    if password_hash and verify_password(password, password_hash):
        return user, "ok"

    legacy_password = user.get("password")
    if legacy_password and legacy_password == password:
        migrate_legacy_password(user["id"], password)
        return user, "ok"

    if not password_hash and not legacy_password:
        return None, "missing_credentials"

    return None, "invalid_password"


def fetch_user_papers(user_id, columns="*"):
    return (
        supabase.table("papers")
        .select(columns)
        .eq("user_id", user_id)
        .order("display_order")
        .execute()
    )


def sort_papers_dataframe(df, sort_option):
    if df.empty:
        return df

    sorted_df = df.copy()

    if sort_option == "年（新しい順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=False)
    elif sort_option == "年（古い順）":
        sorted_df = sorted_df.sort_values(by="year", ascending=True)
    elif sort_option == "タイトル":
        sorted_df = sorted_df.sort_values(by="title", ascending=True)
    elif sort_option == "ステータス":
        sorted_df = sorted_df.sort_values(by="status", ascending=True)

    sorted_df = sorted_df.reset_index(drop=True)
    sorted_df["ref_no"] = sorted_df.index + 1
    return sorted_df


def make_word_citation(row, style="APA"):
    authors = row.get("authors", "")
    year = row.get("year", "")
    title = row.get("title", "")
    journal = row.get("journal", "")
    doi = row.get("doi", "")

    if style == "APA":
        citation = f"{authors} ({year}). {title}. {journal}."
        if doi:
            citation += f" https://doi.org/{doi}"
    elif style == "Vancouver":
        citation = f"{authors}. {title}. {journal}. {year}."
        if doi:
            citation += f" doi:{doi}"
    elif style == "Nature":
        citation = f"{authors} {title}. {journal} ({year})."
        if doi:
            citation += f" https://doi.org/{doi}"
    else:
        citation = f"{authors} ({year}). {title}. {journal}."

    return citation


# -----------------------------
# Word出力
# -----------------------------
def export_to_word_bytes(papers):
    doc = Document()
    doc.add_heading("参考文献", 0)

    for index, paper in enumerate(papers, start=1):
        text = (
            f"[{index}] {paper.get('authors', '')} ({paper.get('year', '')}). "
            f"{paper.get('title', '')}. {paper.get('journal', '')}."
        )
        doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------
# DOI取得
# -----------------------------
def fetch_doi(doi):
    normalized_doi = normalize_doi(doi)
    if not normalized_doi:
        return None

    try:
        response = requests.get(
            f"https://api.crossref.org/works/{normalized_doi}",
            timeout=15,
        )
        response.raise_for_status()
    except requests.RequestException:
        return None

    data = response.json().get("message", {})
    title = data["title"][0] if data.get("title") else ""
    authors = ", ".join(author.get("family", "") for author in data.get("author", []))
    journal = data["container-title"][0] if data.get("container-title") else ""

    issued = data.get("issued", {}).get("date-parts", [])
    year = issued[0][0] if issued and issued[0] else 0

    return title, authors, journal, year


# -----------------------------
# Storage: PDFアップロード
# -----------------------------
def upload_pdf_to_storage(pdf_file, user_id):
    admin_client = get_supabase_admin()
    if admin_client is None:
        raise RuntimeError("PDF storage upload requires service role configuration")

    filename = pdf_file.name
    name, ext = os.path.splitext(filename)
    safe_name = f"{name}_{int(time.time())}{ext}"
    storage_path = f"{user_id}/{safe_name}"

    admin_client.storage.from_(BUCKET_NAME).upload(
        path=storage_path,
        file=pdf_file.read(),
        file_options={"content-type": "application/pdf"},
    )
    return storage_path


def create_pdf_signed_url(storage_path, expires_in=3600):
    if not isinstance(storage_path, str) or not storage_path.strip():
        return None

    admin_client = get_supabase_admin()
    if admin_client is None:
        return None

    response = (
        admin_client.storage.from_(BUCKET_NAME).create_signed_url(
            storage_path,
            expires_in,
        )
    )

    if isinstance(response, dict):
        return response.get("signedURL") or response.get("signedUrl")
    return None


def delete_pdf_from_storage(storage_path):
    if storage_path:
        admin_client = get_supabase_admin()
        if admin_client is None:
            return
        admin_client.storage.from_(BUCKET_NAME).remove([storage_path])


# -----------------------------
# タグ
# -----------------------------
def get_or_create_tag_id(tag_name):
    tag_result = (
        supabase.table("tags")
        .select("id")
        .eq("name", tag_name)
        .limit(1)
        .execute()
    )

    if tag_result.data:
        return tag_result.data[0]["id"]

    new_tag = supabase.table("tags").insert({"name": tag_name}).execute()
    return new_tag.data[0]["id"]


def save_tags_for_paper(paper_id, tags_text):
    for tag_name in normalize_tag_input(tags_text):
        tag_id = get_or_create_tag_id(tag_name)
        supabase.table("paper_tags").upsert(
            {"paper_id": paper_id, "tag_id": tag_id}
        ).execute()


def get_tag_map_for_papers(paper_ids):
    if not paper_ids:
        return {}

    paper_tag_result = (
        supabase.table("paper_tags")
        .select("paper_id, tag_id")
        .in_("paper_id", paper_ids)
        .execute()
    )

    paper_tags = paper_tag_result.data or []
    if not paper_tags:
        return {}

    tag_ids = sorted({row["tag_id"] for row in paper_tags})
    tag_result = supabase.table("tags").select("id, name").in_("id", tag_ids).execute()
    tag_name_map = {row["id"]: row["name"] for row in (tag_result.data or [])}

    tag_map = {paper_id: [] for paper_id in paper_ids}
    for row in paper_tags:
        tag_name = tag_name_map.get(row["tag_id"])
        if tag_name:
            tag_map.setdefault(row["paper_id"], []).append(tag_name)

    return tag_map


# -----------------------------
# 並び順・更新・削除
# -----------------------------
def move_paper(user_id, paper_id, display_order, direction):
    operator = "lt" if direction == "up" else "gt"
    descending = direction == "up"

    neighbor_result = (
        getattr(
            supabase.table("papers")
            .select("id, display_order")
            .eq("user_id", user_id),
            operator,
        )("display_order", display_order)
        .order("display_order", desc=descending)
        .limit(1)
        .execute()
    )

    if not neighbor_result.data:
        return

    neighbor = neighbor_result.data[0]

    (
        supabase.table("papers")
        .update({"display_order": neighbor["display_order"]})
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )

    (
        supabase.table("papers")
        .update({"display_order": display_order})
        .eq("id", neighbor["id"])
        .eq("user_id", user_id)
        .execute()
    )


def update_paper_details(user_id, paper_id, status, notes):
    (
        supabase.table("papers")
        .update({"status": status, "notes": notes})
        .eq("id", paper_id)
        .eq("user_id", user_id)
        .execute()
    )


def delete_paper(user_id, row):
    pdf_path = row.get("pdf_path")
    if isinstance(pdf_path, str) and pdf_path.strip():
        delete_pdf_from_storage(pdf_path)

    supabase.table("paper_tags").delete().eq("paper_id", row["id"]).execute()
    (
        supabase.table("papers")
        .delete()
        .eq("id", row["id"])
        .eq("user_id", user_id)
        .execute()
    )


# -----------------------------
# 認証
# -----------------------------
if "user_id" not in st.session_state:
    st.title("ログイン")

    auth_mode = st.radio("選択", ["ログイン", "新規登録"])

    with st.form("auth_form"):
        username = st.text_input("ユーザー名")
        password = st.text_input("パスワード", type="password")
        submit_label = "登録" if auth_mode == "新規登録" else "ログイン"
        submitted = st.form_submit_button(submit_label)

    if auth_mode == "新規登録":
        if submitted:
            try:
                create_user_with_password_hash(normalize_username(username), password)
                st.success("登録完了")
            except APIError:
                st.error("ユーザー名が既に存在するか、DB設定が不足している可能性があります")
            except Exception:
                st.error("ユーザー名が既に存在する可能性があります")
    else:
        if submitted:
            user, auth_status = authenticate_user(username, password)

            if user:
                st.session_state["user_id"] = user["id"]
                st.session_state["username"] = user["username"]
                st.success("ログイン成功")
                st.rerun()
            elif auth_status == "missing_credentials":
                st.error("このアカウントの認証データが見つかりません。移行前に旧パスワードが削除された可能性があります。")
            else:
                st.error("失敗")

    admin_recovery_token = get_admin_recovery_token()
    if admin_recovery_token and has_service_role_access():
        with st.expander("管理者用パスワード復旧"):
            st.caption("管理トークンを使って、既存ユーザーのパスワードを再設定します。")
            with st.form("admin_recovery_form"):
                recovery_token = st.text_input("管理トークン", type="password")
                recovery_username = st.text_input("復旧するユーザー名")
                recovery_password = st.text_input("新しいパスワード", type="password")
                recovery_submitted = st.form_submit_button("パスワードを再設定")

            if recovery_submitted:
                if recovery_token != admin_recovery_token:
                    st.error("管理トークンが違います。")
                elif not normalize_username(recovery_username) or not recovery_password:
                    st.error("ユーザー名と新しいパスワードを入力してください。")
                else:
                    try:
                        result = reset_user_password_by_admin(
                            recovery_username,
                            recovery_password,
                        )
                        if result.data:
                            st.success("パスワードを再設定しました。")
                        else:
                            st.error("対象ユーザーが見つかりませんでした。")
                    except Exception as error:
                        st.error(f"復旧に失敗しました: {error}")
    elif admin_recovery_token and not has_service_role_access():
        st.warning("管理者復旧トークンはありますが、service role key が未設定のため復旧機能は無効です。")

    st.stop()


# -----------------------------
# アプリ共通UI
# -----------------------------
if st.sidebar.button("ログアウト"):
    st.session_state.clear()
    st.rerun()

st.sidebar.write(f"ログイン中: {st.session_state.get('username', '')}")

st.title("📚 文献管理アプリ")
menu = st.sidebar.selectbox("メニュー", ["追加", "検索", "一覧", "タグ検索"])


# -----------------------------
# 文献追加
# -----------------------------
if menu == "追加":
    user_id = get_current_user_id()
    st.header("文献追加")
    pdf_feature_enabled = has_service_role_access()

    title = st.text_input("タイトル", value=st.session_state.get("title", ""))
    authors = st.text_input("著者", value=st.session_state.get("authors", ""))
    journal = st.text_input("雑誌", value=st.session_state.get("journal", ""))
    year = st.number_input("年", value=int(st.session_state.get("year", 2024)), step=1)
    pdf_file = st.file_uploader("PDFアップロード", type=["pdf"])
    if not pdf_feature_enabled:
        st.caption("PDF保存は現在無効です。service role key を設定すると利用できます。")
    doi = st.text_input("DOI")

    if st.button("DOIから自動入力"):
        result = fetch_doi(doi)
        if result:
            for field_name, value in zip(DOI_FORM_FIELDS, result):
                st.session_state[field_name] = value
            st.rerun()
        st.error("取得失敗")

    tags = st.text_input("タグ（カンマ区切り）")
    status = st.selectbox("読書ステータス", READING_STATUSES)
    notes = st.text_area("抄録メモ", height=150)

    if st.button("追加"):
        normalized_doi = normalize_doi(doi)

        if normalized_doi:
            existing = (
                supabase.table("papers")
                .select("id, title")
                .eq("user_id", user_id)
                .eq("doi", normalized_doi)
                .limit(1)
                .execute()
            )

            if existing.data:
                st.warning("このDOIの文献はすでに登録されています。")
                st.stop()

        try:
            pdf_path = None
            if pdf_file:
                if not pdf_feature_enabled:
                    st.error("PDF保存は現在利用できません。")
                    st.stop()
                pdf_path = upload_pdf_to_storage(pdf_file, user_id)

            max_result = (
                supabase.table("papers")
                .select("display_order")
                .eq("user_id", user_id)
                .order("display_order", desc=True)
                .limit(1)
                .execute()
            )
            current_max = max_result.data[0]["display_order"] if max_result.data else 0
            next_order = (current_max or 0) + 1

            insert_result = (
                supabase.table("papers")
                .insert(
                    {
                        "title": title,
                        "authors": authors,
                        "journal": journal,
                        "year": int(year),
                        "doi": normalized_doi or None,
                        "pdf_path": pdf_path,
                        "user_id": user_id,
                        "display_order": next_order,
                        "status": status,
                        "notes": notes,
                    }
                )
                .execute()
            )

            paper_id = insert_result.data[0]["id"]
            save_tags_for_paper(paper_id, tags)
            st.success("追加しました！")
        except Exception as error:
            st.error(f"エラー内容: {error}")
            st.exception(error)


# -----------------------------
# 検索
# -----------------------------
elif menu == "検索":
    user_id = get_current_user_id()
    keyword = st.text_input("キーワード").strip()

    if st.button("検索"):
        result = fetch_user_papers(user_id, "id, title, authors, year")
        papers = result.data or []

        if keyword:
            lowered_keyword = keyword.lower()
            papers = [
                paper
                for paper in papers
                if lowered_keyword in (paper.get("title") or "").lower()
                or lowered_keyword in (paper.get("authors") or "").lower()
            ]

        if not papers:
            st.write("見つかりません")
        else:
            for paper in papers:
                st.write((paper["id"], paper["title"], paper["authors"], paper["year"]))


# -----------------------------
# 一覧
# -----------------------------
elif menu == "一覧":
    user_id = get_current_user_id()
    result = fetch_user_papers(user_id)
    df = pd.DataFrame(result.data or [])

    sort_option = st.selectbox("並び替え", SORT_OPTIONS)
    df = sort_papers_dataframe(df, sort_option)

    st.header("📚 論文一覧")

    if not df.empty:
        word_bytes = export_to_word_bytes(df.to_dict(orient="records"))
        st.download_button(
            "📄 Word出力",
            data=word_bytes,
            file_name="references.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if df.empty:
        st.write("データがありません")
    else:
        tag_map = get_tag_map_for_papers(df["id"].tolist())

        for _, row in df.iterrows():
            row_dict = row.to_dict()
            pdf_path = row_dict.get("pdf_path")
            signed_url = create_pdf_signed_url(pdf_path, 3600)

            with st.container():
                st.markdown(f"### [{row_dict['ref_no']}] {row_dict['title']}")
                st.write(f"著者: {row_dict['authors']}")
                st.write(f"雑誌: {row_dict['journal']} ({row_dict['year']})")

                if row_dict.get("status"):
                    st.write(f"ステータス: {row_dict['status']}")

                if row_dict.get("notes"):
                    st.write("メモ:")
                    st.write(row_dict["notes"])

                tags_list = tag_map.get(row_dict["id"], [])
                if tags_list:
                    st.write("タグ:", ", ".join(tags_list))

                col1, col2, col3, col4, col5, col6 = st.columns(6)

                with col1:
                    if signed_url:
                        st.link_button("📄 PDF", signed_url)

                with col2:
                    if signed_url:
                        st.link_button("👀 開く", signed_url)

                with col3:
                    if st.button("🗑 削除", key=f"del_{row_dict['id']}"):
                        delete_paper(user_id, row_dict)
                        st.success("削除しました")
                        st.rerun()

                with col4:
                    if st.button("📚 引用", key=f"cite_{row_dict['id']}"):
                        st.code(make_word_citation(row_dict, style="APA"))

                with col5:
                    if st.button("⬆", key=f"up_{row_dict['id']}"):
                        move_paper(user_id, row_dict["id"], row_dict["display_order"], "up")
                        st.rerun()

                with col6:
                    if st.button("⬇", key=f"down_{row_dict['id']}"):
                        move_paper(
                            user_id,
                            row_dict["id"],
                            row_dict["display_order"],
                            "down",
                        )
                        st.rerun()

                with st.expander("編集"):
                    current_status = row_dict.get("status")
                    status_index = (
                        READING_STATUSES.index(current_status)
                        if current_status in READING_STATUSES
                        else 0
                    )
                    edit_status = st.selectbox(
                        "読書ステータス",
                        READING_STATUSES,
                        index=status_index,
                        key=f"status_{row_dict['id']}",
                    )

                    edit_notes = st.text_area(
                        "抄録メモ",
                        value=row_dict.get("notes") or "",
                        height=150,
                        key=f"notes_{row_dict['id']}",
                    )

                    if st.button("💾 保存", key=f"save_{row_dict['id']}"):
                        try:
                            update_paper_details(
                                user_id,
                                row_dict["id"],
                                edit_status,
                                edit_notes,
                            )
                            st.success("更新しました")
                            st.rerun()
                        except Exception as error:
                            st.error(f"更新失敗: {error}")

                st.divider()


# -----------------------------
# タグ検索
# -----------------------------
elif menu == "タグ検索":
    user_id = get_current_user_id()
    tag = st.text_input("タグ名").strip()

    if st.button("検索"):
        tag_result = (
            supabase.table("tags")
            .select("id")
            .eq("name", tag)
            .limit(1)
            .execute()
        )

        if not tag_result.data:
            st.write("見つかりません")
        else:
            paper_tag_result = (
                supabase.table("paper_tags")
                .select("paper_id")
                .eq("tag_id", tag_result.data[0]["id"])
                .execute()
            )
            paper_ids = [row["paper_id"] for row in (paper_tag_result.data or [])]

            if not paper_ids:
                st.write("見つかりません")
            else:
                papers_result = (
                    supabase.table("papers")
                    .select("id, title")
                    .eq("user_id", user_id)
                    .in_("id", paper_ids)
                    .execute()
                )

                if not papers_result.data:
                    st.write("見つかりません")
                else:
                    for paper in papers_result.data:
                        st.write((paper["id"], paper["title"]))
